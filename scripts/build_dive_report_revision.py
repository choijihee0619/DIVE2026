from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


BASE = Path("/Users/choijihee/Library/CloudStorage/GoogleDrive-heeppiness@pukyong.ac.kr/내 드라이브/DIVE2026")
SRC = BASE / "개발설계보고서_260714.docx"
OUT = BASE / "개발설계보고서_260714_수정보완.docx"


table_no = 0
fig_no = 0


def east_asia(run, font="맑은 고딕"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def cell_text(cell, text, bold=False, size=8.0, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(str(text))
    east_asia(r)
    r.font.size = Pt(max(size, 6.6))
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(begin)
    r._r.append(instr)
    r._r.append(end)


def add_toc(paragraph):
    r = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "목차는 Word에서 문서 열기 시 필드 업데이트로 갱신됩니다."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(begin)
    r._r.append(instr)
    r._r.append(sep)
    r._r.append(text)
    r._r.append(end)


def caption(doc, text, kind="표"):
    global table_no, fig_no
    if kind == "표":
        table_no += 1
        label = f"표 {table_no}. {text}"
    else:
        fig_no += 1
        label = f"그림 {fig_no}. {text}"
    p = doc.add_paragraph(label)
    p.style = doc.styles["Caption"]
    return p


def table(doc, title, headers, rows, widths=None, size=7.2):
    caption(doc, title)
    t = doc.add_table(rows=1, cols=len(headers))
    t.autofit = False
    hdr = t.rows[0]
    repeat_header(hdr)
    for i, h in enumerate(headers):
        cell_text(hdr.cells[i], h, bold=True, size=size, color="0B2545")
        set_cell_shading(hdr.cells[i], "E8EEF5")
        if widths:
            set_cell_width(hdr.cells[i], widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cell_text(cells[i], v, size=size)
            if widths:
                set_cell_width(cells[i], widths[i])
    return t


def kv(doc, title, pairs):
    caption(doc, title)
    t = doc.add_table(rows=0, cols=2)
    t.autofit = False
    for k, v in pairs:
        row = t.add_row().cells
        cell_text(row[0], k, bold=True, size=8.2, color="0B2545")
        set_cell_shading(row[0], "F2F4F7")
        cell_text(row[1], v, size=8.2)
        set_cell_width(row[0], 4.2)
        set_cell_width(row[1], 12.4)
    return t


def para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    east_asia(r)
    return p


def bullet(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        east_asia(r)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        east_asia(r)


def h1(doc, text):
    doc.add_heading(text, 1)


def h2(doc, text):
    doc.add_heading(text, 2)


def h3(doc, text):
    doc.add_heading(text, 3)


def code_lines(doc, title, lines):
    global fig_no
    caption(doc, title, kind="그림")
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        r.font.size = Pt(8)


def ref_box(doc, title, items):
    kv(doc, title, [("상세 문서 참조", "\n".join(items))])


def style_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Pt(595.3)
    sec.page_height = Pt(841.9)
    sec.top_margin = Pt(62)
    sec.bottom_margin = Pt(62)
    sec.left_margin = Pt(60)
    sec.right_margin = Pt(60)
    sec.header_distance = Pt(34)
    sec.footer_distance = Pt(34)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.1
    normal.paragraph_format.space_after = Pt(4)
    for name, size, color, before, after in [
        ("Title", 22, "0B2545", 0, 10),
        ("Subtitle", 11, "334155", 0, 8),
        ("Heading 1", 14, "0B2545", 13, 6),
        ("Heading 2", 11.5, "1F4D78", 8, 4),
        ("Heading 3", 10, "334155", 6, 3),
        ("Caption", 8.2, "64748B", 4, 2),
    ]:
        st = styles[name]
        st.font.name = "맑은 고딕"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = name.startswith("Heading") or name == "Title"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = name.startswith("Heading")
    for section in doc.sections:
        hp = section.header.paragraphs[0]
        hp.text = "HUG × 아이엔 안심주거 생태계 개발설계보고서"
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp = section.footer.paragraphs[0]
        fp.text = "DIVE 2026 | 2026.07.14 | "
        add_page_number(fp)


DOCS = [
    ["개발설계보고서_260714.docx", "프로젝트 전체 기획과 문서 간 연결 기준", "기획자, 전체 개발자, 발표자", "비전, 생애주기, 상위 아키텍처, MVP, 공통 식별자·상태", "API 파라미터, 전체 ERD, 전체 화면정의, Solidity 상세", "없음", "5개 상세 명세서와 구현 로드맵", "DOCX"],
    ["데이터수집_및_API가이드_260714.md", "데이터·외부 API 수집과 정규화 기준", "백엔드, 데이터, ML", "발제사 데이터, CODEF, 공공 API, Mock, 인증·비용·제약, 요청·응답", "UI 컴포넌트, 스마트컨트랙트 함수", "개발설계보고서", "API Adapter, Mock Dataset", "Markdown"],
    ["ML개발가이드_260714.md", "Colab 기반 모델 학습·평가·배포 기준", "ML, 백엔드", "Notebook, 전처리, Feature, 모델, SHAP, 저장·버전", "REST API 상세, 화면 Wireframe", "개발설계보고서, 데이터수집 가이드", "ipynb, model artifact, model card", "Markdown"],
    ["Frontend_UIUX_명세서_260714.md", "화면·상태·컴포넌트 구현 기준", "프론트엔드, 디자이너", "IA, Sitemap, User Flow, Wireframe, 페이지 정의, 상태 UI", "DB 컬럼, API 인증 절차", "개발설계보고서", "Next.js/React 화면 코드", "Markdown"],
    ["Backend_API_명세서_260714.md", "FastAPI·DB·스케줄러·권한 구현 기준", "백엔드, 프론트엔드, ML", "REST API, Adapter, ERD, 테이블, 오류코드, 인증·권한, 배포", "ML 학습 코드, Solidity 코드", "개발설계보고서, 데이터수집 가이드", "FastAPI 서버, DB migration", "Markdown"],
    ["Blockchain_설계서_260714.md", "온체인 감사장부와 스마트컨트랙트 구현 기준", "블록체인, 백엔드", "Solidity, 상태머신, 역할, 이벤트, 배포, 테스트, 개인정보 보호", "외부 API 요청·응답, UI 상세", "개발설계보고서", "Hardhat project, contract ABI", "Markdown"],
]


MODULES = [
    ["Frontend", "사용자별 화면, 입력 검증, 상태 표시", "사용자 입력, API 응답", "화면 상태, CTA", "Backend API", "오류 컴포넌트와 재시도", "MVP", "Frontend_UIUX_명세서_260714.md"],
    ["Backend API", "인증, 라우팅, 요청 검증, 서비스 오케스트레이션", "REST 요청", "표준 응답, 오류코드", "모든 서비스", "공통 오류 응답", "MVP", "Backend_API_명세서_260714.md"],
    ["Contract Service", "계약 생성, 상태, 버전 연결", "contract_id, 계약정보", "ContractStatus, 버전 참조", "Property, Timeline", "상태 전이 거부", "MVP", "Backend_API_명세서_260714.md"],
    ["Property Service", "주소 정규화와 물건 기준키 관리", "주소, 동·호", "property_id, 주소키", "External API Gateway", "수동 후보 선택", "MVP", "데이터수집_및_API가이드_260714.md"],
    ["Risk Engine", "규칙 기반 위험 판단과 등급 산정", "등기·시세·계약정보", "risk_assessment_id, 등급", "ML, RAG", "RuleOnlyFallback", "MVP", "Backend_API_명세서_260714.md"],
    ["ML Inference Service", "위험유사도·회수등급·처리기간 추론", "feature payload, model_version_id", "모델 결과, SHAP", "Model Artifact", "InsufficientData", "PoC/MVP 일부", "ML개발가이드_260714.md"],
    ["RAG Service", "상담·제도 근거 검색", "질문, 메타데이터", "근거 문서, 요약", "Vector Store, LLM", "근거 없음 표시", "MVP", "ML개발가이드_260714.md"],
    ["Evidence Service", "증빙 요청·업로드·해시", "파일, request_id", "evidence_id, document_hash", "Object Storage", "추출 실패", "MVP", "Backend_API_명세서_260714.md"],
    ["Verification Service", "증빙·API 결과 검토와 상태 확정", "evidence_id, 검증자", "verification_id, VerificationStatus", "Blockchain Adapter", "Rejected/Expired", "MVP", "Backend_API_명세서_260714.md"],
    ["Timeline Service", "계약 생애주기 이벤트 통합", "상태 변경", "timeline_event_id", "Contract, Blockchain", "중복 이벤트 방지", "MVP", "Backend_API_명세서_260714.md"],
    ["Scheduler Service", "D-90 대상 조회와 작업 생성", "계약 만료일", "return_plan 요청 이벤트", "Notification, Timeline", "재시도 큐", "MVP", "Backend_API_명세서_260714.md"],
    ["Notification Service", "이메일/SMS/카카오/FCM 알림", "수신자, 템플릿", "발송 결과", "Scheduler", "대체 채널", "MVP 일부", "Backend_API_명세서_260714.md"],
    ["Incident Service", "미반환 사고 접수와 HUG 인계", "사고 신고", "incident_id, 상태", "Recovery, Timeline", "중복 접수 방지", "MVP", "Backend_API_명세서_260714.md"],
    ["Recovery Service", "회수등급·처리기간 기반 우선순위", "incident_id", "recovery_prediction_id", "ML", "수동 우선순위", "PoC", "ML개발가이드_260714.md"],
    ["External API Gateway", "CODEF·공공 API 호출·로그·fallback", "정규화 주소, 요청값", "스냅샷, api_call_id", "Provider APIs", "MockFallback", "MVP", "데이터수집_및_API가이드_260714.md"],
    ["Blockchain Adapter", "해시·상태 이벤트 트랜잭션", "참조 ID, 해시", "blockchain_tx_id", "Polygon Amoy", "DB 감사로그 대체", "MVP", "Blockchain_설계서_260714.md"],
    ["PostgreSQL", "관계형 원장, 상태, 업무 데이터", "서비스 쓰기", "영속 데이터", "Backend", "트랜잭션 rollback", "MVP", "Backend_API_명세서_260714.md"],
    ["Object Storage", "계약서·등기부·증빙 원문 저장", "파일", "object_uri, hash", "Evidence", "재업로드 요청", "MVP", "Backend_API_명세서_260714.md"],
    ["Vector Store", "상담·제도 문서 임베딩 검색", "embedding", "유사 문서", "RAG", "키워드 검색 대체", "MVP", "ML개발가이드_260714.md"],
]


IDS = [
    ["user_id", "Backend/Auth", "사용자와 역할 식별", "Backend_API_명세서_260714.md"],
    ["property_id", "Property Service", "주소 정규화 후 물건 식별", "Backend_API_명세서_260714.md"],
    ["contract_id", "Contract Service", "계약 생애주기 중심 ID", "Backend_API_명세서_260714.md"],
    ["contract_version_id", "Contract Service", "계약서 버전 식별", "Backend_API_명세서_260714.md"],
    ["registry_snapshot_id", "External API Gateway", "등기/API 조회 시점 스냅샷", "데이터수집_및_API가이드_260714.md"],
    ["risk_assessment_id", "Risk Engine", "위험진단 결과 식별", "Backend_API_명세서_260714.md"],
    ["evidence_request_id", "Evidence Service", "증빙 요청 식별", "Backend_API_명세서_260714.md"],
    ["evidence_id", "Evidence Service", "제출 증빙 식별", "Backend_API_명세서_260714.md"],
    ["verification_id", "Verification Service", "검증 결과 식별", "Backend_API_명세서_260714.md"],
    ["return_plan_id", "Scheduler/Return Plan", "D-90 반환계획 식별", "Backend_API_명세서_260714.md"],
    ["incident_id", "Incident Service", "사고 접수 식별", "Backend_API_명세서_260714.md"],
    ["recovery_prediction_id", "Recovery Service", "회수예측 결과 식별", "ML개발가이드_260714.md"],
    ["timeline_event_id", "Timeline Service", "상태 이벤트 식별", "Backend_API_명세서_260714.md"],
    ["model_version_id", "ML Pipeline", "모델 아티팩트 버전 식별", "ML개발가이드_260714.md"],
    ["blockchain_tx_id", "Blockchain Adapter", "트랜잭션 로그 식별", "Blockchain_설계서_260714.md"],
    ["api_call_id", "External API Gateway", "외부 API 호출 로그 식별", "데이터수집_및_API가이드_260714.md"],
]


def build():
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    doc = Document()
    style_doc(doc)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("HUG × 아이엔 안심주거 생태계\n개발기획 및 상세설계보고서")
    sp = doc.add_paragraph(style="Subtitle")
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.add_run("6개 문서 체계에 맞춘 최상위 개발 기준 문서")
    kv(doc, "표지 정보", [
        ("프로젝트명", "HUG 안심전세 체인"),
        ("작성 기준일", "2026년 7월 14일(KST)"),
        ("문서 역할", "프로젝트 전체 구조, 문서 간 연결, 공통 식별자와 상태, MVP 범위, 개발 순서를 통제하는 최상위 문서"),
        ("핵심 메시지", "위험을 점수화하는 앱이 아니라, 계약 위험이 실제로 해소됐는지 증명하고 사고 전후의 대응을 끊김 없이 연결하는 주거안전 인프라"),
        ("상세 문서", "데이터수집/API, ML, Frontend UI/UX, Backend API, Blockchain 설계서는 별도 Markdown으로 관리"),
    ])
    doc.add_page_break()

    h1(doc, "문서 변경이력")
    table(doc, "변경이력", ["버전", "작성일", "변경사항"], [[
        "1.1", "2026.07.14",
        "개발설계보고서를 6개 문서 체계의 최상위 기준 문서로 재구성. API·ML·UI·DB·Solidity 상세는 별도 명세서로 이관하고, 문서 구조도, 모듈 책임, 공통 식별자, 공통 상태, 상위 아키텍처, 개발 순서를 보강."
    ]], [2.0, 3.0, 11.5], 8.2)

    h1(doc, "문서 체계 및 상세 명세서 목록")
    code_lines(doc, "문서 구조도", [
        "개발설계보고서",
        "├── 데이터수집 및 API 가이드",
        "├── ML 개발가이드",
        "├── Frontend UI/UX 명세서",
        "├── Backend API 명세서",
        "└── Blockchain 설계서",
    ])
    table(doc, "문서별 책임과 경계", ["문서명", "문서 목적", "주요 독자", "포함 내용", "포함하지 않는 내용", "선행 문서", "후속 산출물", "형식"], DOCS, [2.6, 3.0, 2.3, 3.6, 3.2, 2.0, 2.6, 1.1], 5.8)
    para(doc, "개발설계보고서는 무엇을 만들고 왜 필요한지, 전체 시스템에서 어떤 역할을 하는지, 모듈 간 어떻게 연결되는지를 정의한다. 상세 명세서는 실제 구현에 필요한 필드, 함수, API, 컴포넌트, 컬럼, 테스트와 배포 절차를 정의한다.")
    h2(doc, "문서 간 참조 규칙")
    table(doc, "문서 간 참조 규칙", ["규칙", "적용 기준", "충돌 시 우선순위"], [
        ["공통 식별자", "모든 문서는 본 개발설계보고서 10장의 식별자를 그대로 사용한다.", "개발설계보고서 우선"],
        ["상태값", "ContractStatus, VerificationStatus, BlockchainStatus, APIResultStatus, ModelResultStatus는 본 문서의 명칭을 기준으로 한다.", "개발설계보고서 우선, 전이는 상세 문서"],
        ["외부 API", "Endpoint, 파라미터, 응답 JSON, 인증, 비용은 데이터수집 가이드에서만 상세 정의한다.", "데이터수집_및_API가이드 우선"],
        ["내부 REST API", "경로, 요청/응답 DTO, 오류코드는 Backend API 명세서에서만 상세 정의한다.", "Backend_API_명세서 우선"],
        ["UI 상태", "화면별 Default/Loading/Error 등 세부 상태는 Frontend 명세서에서 정의한다.", "Frontend_UIUX_명세서 우선"],
        ["ML Feature", "Feature Engineering, 학습/검증 방식, 모델 저장은 ML 개발가이드에서 정의한다.", "ML개발가이드 우선"],
        ["온체인 이벤트", "Solidity 함수, 이벤트, role, 배포 주소는 Blockchain 설계서에서 정의한다.", "Blockchain_설계서 우선"],
        ["문서 업데이트", "상세 명세 변경이 공통 식별자·상태에 영향을 주면 개발설계보고서를 먼저 갱신한다.", "개발설계보고서 갱신 후 상세 문서 반영"],
    ], [3.0, 8.0, 5.5], 7.0)
    h2(doc, "AI 코드 생성용 문서 사용 원칙")
    bullet(doc, [
        "UI 코드를 생성할 때는 개발설계보고서의 상위 IA를 확인한 뒤 Frontend_UIUX_명세서_260714.md의 화면 정의를 사용한다.",
        "백엔드 코드를 생성할 때는 공통 식별자와 상태를 먼저 확인하고 Backend_API_명세서_260714.md의 DTO·ERD·오류코드를 사용한다.",
        "ML 코드를 생성할 때는 데이터수집 가이드의 데이터 정의와 ML개발가이드의 Notebook 구조를 함께 참조한다.",
        "블록체인 코드를 생성할 때는 개발설계보고서의 온체인 금지 원칙을 먼저 적용하고 Blockchain_설계서_260714.md의 Solidity 명세를 사용한다.",
    ])
    doc.add_page_break()

    h1(doc, "목차")
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    h1(doc, "0. 문서 개요")
    h2(doc, "0.1 문서 목적")
    para(doc, "본 문서는 HUG 안심전세 체인의 최상위 개발 기준 문서다. 프로젝트 비전, 발제 요구사항 해석, 서비스 생애주기, 상위 아키텍처, 모듈 책임, 공통 식별자·상태, MVP 범위, 통합 순서와 상세 문서 참조 관계를 정의한다.")
    h2(doc, "0.2 프로젝트 범위")
    para(doc, "프로젝트는 계약 전 위험진단에서 끝나지 않고 위험조건 해소 여부를 검증하며, 계약 만료 전 반환 준비와 사고 후 HUG 채권관리까지 연결하는 전세 계약 생애주기 플랫폼이다.")
    h2(doc, "0.3 문서 책임")
    bullet(doc, ["개발설계보고서: 전체 구조와 공통 기준", "상세 명세서: 실제 구현 기준", "중복 작성 금지: 동일한 세부 필드·함수·화면·API는 하나의 책임 문서에만 상세 정의"])
    h2(doc, "0.4 상세 문서 책임")
    table(doc, "상세 문서 책임 요약", ["영역", "개발설계보고서에 남기는 것", "상세 문서로 이관하는 것"], [
        ["데이터/API", "데이터 역할, 상위 흐름, MVP 구분", "인증, Endpoint, 파라미터, 응답, 저장, 전처리"],
        ["ML", "모델 목적, 입출력, 시스템 역할", "Notebook, Feature, 학습, 검증, 아티팩트"],
        ["Frontend", "상위 IA, 핵심 화면, User Flow", "페이지별 상태, 컴포넌트, Wireframe, 반응형"],
        ["Backend", "서비스 구조, 핵심 엔터티, 상위 API", "REST API, ERD, 테이블, 오류코드, 인증"],
        ["Blockchain", "목적, 기록 대상, 시스템 연결", "Solidity, 함수, 이벤트, 역할, 테스트, 배포"],
    ], [3.0, 6.6, 6.9], 7.4)
    h2(doc, "0.5 용어 및 표기 기준")
    bullet(doc, ["식별자는 snake_case를 사용한다.", "상태값은 PascalCase를 사용한다.", "일시는 KST 기준 ISO 8601을 우선한다.", "금액은 KRW 정수 원 단위로 저장하고 화면에서 천 단위 구분을 적용한다.", "주소는 도로명주소, 지번주소, 법정동코드, 건물관리번호를 구분한다."])
    ref_box(doc, "0장 상세 문서 참조", ["문서 책임 경계: 본 개발설계보고서", "각 상세 구현 기준: 해당 Markdown 명세서"])

    h1(doc, "1. 프로젝트 배경과 문제 정의")
    h2(doc, "1.1 발제 요구사항")
    para(doc, "HUG와 아이엔의 발제 맥락은 전세 계약 위험을 조기에 발견하고, 위험 해소 조건이 실제 이행됐는지 검증하며, 사고 발생 후 HUG 채권관리 단계로 데이터가 이어지도록 하는 것이다.")
    h2(doc, "1.2 임대차 보호 구조의 단절")
    bullet(doc, ["계약 전 등기·시세·건축물 정보가 분산되어 임차인이 직접 해석하기 어렵다.", "근저당 말소나 특약 반영 등 보완조건의 실제 이행 여부를 공동으로 확인하기 어렵다.", "계약 만료 전 반환 준비 확인이 늦어 이상징후가 사고 직전까지 누락된다.", "사고 후 HUG 채권관리 단계에서 계약 전 위험정보와 검증 이력이 충분히 연결되지 않는다."])
    h2(doc, "1.3 해결하려는 문제")
    code_lines(doc, "핵심 서비스 흐름", [
        "계약 전 위험진단 -> 위험조건 보완요청 -> 외부 API·증빙 검증",
        "-> 계약서 및 검증결과 블록체인 기록 -> 계약기간 중 상태변화 관리",
        "-> D-90 반환계획 확인 -> 이상징후 조기경보 -> 보증금 미반환 사고",
        "-> HUG 채권관리 인계 -> 예상 회수등급·처리기간 기반 사건 우선순위 보조",
    ])
    h2(doc, "1.4 프로젝트 목표")
    numbered(doc, ["위험을 발견하고 근거를 설명한다.", "해결 가능한 위험과 해결 불가능한 위험을 구분한다.", "위험조건 보완 요청과 증빙 검증을 표준화한다.", "계약 버전과 검증 이력을 감사 가능한 형태로 남긴다.", "D-90 반환계획 확인과 사고 후 채권관리 인계를 연결한다."])
    h2(doc, "1.5 서비스 정의")
    para(doc, "HUG 안심전세 체인은 계약 전 위험진단에서 끝나지 않고, 위험조건이 실제로 해소되었는지 검증하며, 계약 만료 전 반환 준비와 사고 후 HUG 채권관리까지 연결하는 전세 계약 생애주기 플랫폼이다.")
    h2(doc, "1.6 핵심 가치")
    table(doc, "핵심 가치", ["가치", "설명"], [["증명 가능성", "계약서, 등기 조회, 증빙 검증, 상태 변경을 해시와 이벤트로 추적"], ["조기개입", "D-90 반환계획 확인으로 사고 전 대응 시점을 앞당김"], ["업무 연결", "상담·검증·보증·채권관리로 이어지는 데이터 흐름 제공"], ["과장 방지", "사고확률, 법적 집행, 금융 이체 등 검증되지 않은 기능을 명확히 제한"]], [4.0, 12.5])

    h1(doc, "2. 사용자와 이해관계자")
    table(doc, "사용자 정의와 책임", ["사용자", "Pain Point", "기대효과", "권한과 책임"], [
        ["임차인", "등기·시세·계약조건 해석 어려움", "위험근거와 필요한 행동 확인", "계약 입력, 증빙 확인, 사고 접수"],
        ["임대인", "정상 계약 증빙 제출 경로 분산", "보완조건과 제출 상태 명확화", "증빙 제출, 반환계획 제출"],
        ["아이엔 상담·검증 담당자", "상담 분류와 검증 반복 업무", "RAG 근거와 검증 워크플로우", "전문가 이관, 검증 상태 입력"],
        ["HUG 채권관리 담당자", "사고 우선순위 판단과 유사사건 비교 부담", "회수등급·처리기간 보조", "사건 검토, 담당자 액션"],
        ["시스템 관리자", "권한·API·배포·로그 관리", "통합 운영 기준", "사용자 권한, 환경변수, 장애 대응"],
        ["외부 API·검증기관", "조회 결과의 역할과 보관 범위 명확화 필요", "조회 시점과 결과 해시 관리", "API 결과 제공, 검증 상태 참조"],
    ], [2.8, 4.0, 4.5, 5.2], 7.0)

    h1(doc, "3. 서비스 생애주기")
    table(doc, "생애주기 단계", ["단계", "목표", "주요 입력", "주요 출력"], [
        ["계약 전", "주소·계약정보 기반 위험 발견", "주소, 보증금, 임대인 유형", "risk_assessment_id"],
        ["위험조건 보완", "해결 가능한 위험의 조건화", "위험요소, 계약조건", "evidence_request_id"],
        ["검증 및 계약 확정", "증빙·API 결과 대조와 계약 버전 확정", "증빙, 등기 스냅샷, 계약서", "verification_id, contract_version_id"],
        ["계약 중 모니터링", "상태변화와 이력 관리", "계약 상태, 재조회 결과", "timeline_event_id"],
        ["D-90 반환계획", "반환 준비 확인과 이상징후 조기경보", "만료일, 반환계획", "return_plan_id"],
        ["사고 발생", "미반환 의심 접수와 안내", "신고, 계약 이력", "incident_id"],
        ["HUG 채권관리", "회수등급·처리기간 기반 우선순위 보조", "사고·경매·배당 데이터", "recovery_prediction_id"],
        ["계약 종료", "종결 상태와 감사 이력 보존", "종결 사유", "Closed 상태"],
    ], [2.5, 5.2, 4.4, 4.4], 7.0)

    h1(doc, "4. 서비스 범위")
    table(doc, "MVP·Mock·PoC·확장 구분", ["구분", "포함 범위", "제외 또는 유보 범위"], [
        ["MVP", "주소·계약 입력, CODEF 또는 업로드 대체, 위험등급, 보완요청, 검증 이력, 계약 버전 해시, D-90 데모, HUG 대시보드 Mock", "은행 실이체, 실제 HUG 내부망 연동"],
        ["Mock·데모", "Mock API, 샘플 JSON, 정상/근저당/압류 등기부 샘플, 가상 에스크로 상태전이", "Mock 결과를 실제 예측성능으로 주장하지 않음"],
        ["기관 연계형 PoC", "CODEF 실계정, HUG 데이터, 알림 채널, 채권관리 업무 기준 검증", "운영 전환은 별도 승인"],
        ["장기 확장", "비식별 법인키 그래프, 은행 에스크로, 법원·세금·금융기관 연계", "해커톤 단독 구현 범위 아님"],
        ["구현하지 않는 기능", "근저당 말소, 압류, 경매개시, 대위변제 승인, 법원 배당순위 결정", "외부기관 권한 또는 법적 절차"],
    ], [3.0, 7.5, 6.0], 7.0)
    table(doc, "MVP 완료 기준", ["영역", "완료 기준", "검수 포인트"], [
        ["계약 입력", "주소, 보증금, 계약기간, 임대인 유형이 contract_id로 저장된다.", "Validation error와 저장 성공 상태"],
        ["외부 조회", "CODEF 또는 Mock/업로드 대체로 registry_snapshot_id가 생성된다.", "Success, Partial, MockFallback 구분"],
        ["위험진단", "규칙 엔진 기반 위험등급과 근거가 risk_assessment_id로 저장된다.", "점수만 표시하지 않고 근거와 다음 행동 포함"],
        ["증빙·검증", "보완요청, 제출, 검토, 검증 결과가 VerificationStatus로 관리된다.", "Rejected/Expired 처리"],
        ["블록체인", "계약 버전 또는 검증 이벤트 하나 이상이 Polygon Amoy 또는 Mock chain log에 기록된다.", "Pending/Confirmed/Failed 표시"],
        ["D-90", "스케줄러 또는 데모 트리거가 반환계획 요청 이벤트를 생성한다.", "스마트컨트랙트 자동실행으로 표현하지 않음"],
        ["HUG 대시보드", "사고/조기경보/회수등급 Mock 데이터가 우선순위 화면에 표시된다.", "모델 실패 시 수동 또는 RuleOnlyFallback"],
        ["Fallback", "외부 API 실패 시 재시도, 업로드, Mock 전환 중 하나로 시연이 계속된다.", "사용자에게 실패 상태를 명확히 표시"],
    ], [3.0, 8.0, 5.5], 7.0)

    h1(doc, "5. 핵심 기능")
    table(doc, "핵심 기능과 상세 문서", ["기능", "상위 역할", "상세 문서"], [
        ["주소·계약 입력", "진단의 기준 계약과 물건 식별", "Frontend_UIUX_명세서_260714.md, Backend_API_명세서_260714.md"],
        ["외부 데이터 조회", "등기·시세·건축물·법인 정보 수집", "데이터수집_및_API가이드_260714.md"],
        ["위험진단", "규칙·ML·RAG를 결합한 위험근거 생성", "ML개발가이드_260714.md, Backend_API_명세서_260714.md"],
        ["보완조건", "해결 가능한 위험을 증빙 요청으로 변환", "Backend_API_명세서_260714.md"],
        ["증빙·검증", "제출 증빙과 외부 조회 결과 검증", "Backend_API_명세서_260714.md"],
        ["계약 버전", "초안·보완·최종·HUG 제출본 일치 확인", "Blockchain_설계서_260714.md"],
        ["블록체인 감사장부", "계약·검증·상태변경 해시와 이벤트 기록", "Blockchain_설계서_260714.md"],
        ["D-90", "만료 전 반환계획 확인과 알림", "Backend_API_명세서_260714.md"],
        ["사고 대응", "미반환 신고와 HUG 인계", "Backend_API_명세서_260714.md"],
        ["채권관리", "회수등급·처리기간 기반 우선순위 보조", "ML개발가이드_260714.md"],
        ["RAG 상담", "상담사례·제도 근거 검색과 설명 생성", "ML개발가이드_260714.md"],
    ], [3.2, 7.0, 6.3], 7.0)

    doc.add_page_break()
    h1(doc, "6. AI·ML·RAG 상위 설계")
    h2(doc, "6.1 AI 사용 목적")
    code_lines(doc, "AI 처리 흐름", [
        "외부 API·계약정보 -> 법적·업무 규칙 엔진 -> ML 위험유사도 또는 회수예측",
        "-> SHAP 주요 근거 -> RAG 유사 상담사례 검색 -> LLM 기반 사용자 친화적 설명",
        "-> 필요한 증빙과 다음 행동 생성 -> 전문가 이관 여부 판단",
    ])
    bullet(doc, ["규칙 엔진은 명시적 위험 판단을 담당한다.", "ML은 패턴·유사도·회수등급·처리기간 예측을 담당한다.", "SHAP은 모델 결과의 설명을 담당한다.", "RAG는 상담 및 제도 근거 검색을 담당한다.", "LLM은 정보를 재구성하지만 독자적으로 법률 결론을 생성하지 않는다.", "고위험 법률 사건은 전문가에게 이관한다."])
    h2(doc, "6.2~6.7 모델과 책임")
    table(doc, "AI·ML·RAG 모델 요약", ["모델/기능", "해결 문제", "입력 데이터", "출력", "사용 위치", "평가기준", "현재 한계"], [
        ["규칙 엔진", "명시적 위험 판단", "등기, 계약정보, 시세, 건축물", "위험요소, 등급 보정", "위험진단 결과", "정책 검토, 테스트 케이스", "법률자문 아님"],
        ["상담 분류 모델", "분쟁유형·진행단계·전문가 이관", "상담본문, 메타데이터", "유형, 단계, 이관 필요도", "아이엔 상담/RAG", "macro-F1", "라벨 품질 의존"],
        ["회수등급 모델", "예상 배당률 또는 회수등급", "HUG 사고·경매·배당 데이터", "회수등급", "HUG 대시보드", "MAE 또는 F1", "경매 회차 데이터 필요"],
        ["처리기간 모델", "사고 처리 소요기간", "절차일자, 물건특성, 채권규모", "예상 처리기간", "HUG 대시보드", "MAE/RMSE", "절차 기준일 정의 필요"],
        ["RAG", "근거 기반 설명", "상담사례, 제도문서, 위험요소", "근거, 설명, 다음 행동", "임차인/상담자 화면", "근거 적합성", "근거 없으면 답변 제한"],
    ], [2.2, 3.2, 3.0, 2.2, 2.3, 2.0, 2.6], 6.3)
    ref_box(doc, "6장 상세 문서 참조", ["ML 학습·평가·모델 저장: ML개발가이드_260714.md", "RAG 저장소와 임베딩: ML개발가이드_260714.md", "추론 API와 오류코드: Backend_API_명세서_260714.md"])

    h1(doc, "7. 데이터 및 외부 API 상위 설계")
    table(doc, "데이터/API 상위 목록", ["데이터/API명", "프로젝트 내 역할", "구분", "핵심 출력", "실패 시 대체경로", "상세 문서"], [
        ["발제사 상담데이터", "상담 분류와 RAG 근거", "PoC", "분쟁유형, 유사사례", "샘플 상담 데이터", "데이터수집_및_API가이드_260714.md"],
        ["HUG 사고·경매·배당데이터", "회수등급·처리기간 모델", "PoC", "회수등급, 처리기간", "샘플 사건 데이터", "데이터수집_및_API가이드_260714.md"],
        ["CODEF 등기부 API", "조회 시점 기준 등기정보 확인", "MVP", "등기 스냅샷", "PDF 업로드, Mock JSON", "데이터수집_및_API가이드_260714.md"],
        ["도로명주소 API", "주소 정규화와 기준키 생성", "MVP", "법정동코드, 건물관리번호", "수동 후보 선택", "데이터수집_및_API가이드_260714.md"],
        ["국토부 실거래가", "전세가율과 이상계약 보조", "MVP", "주변 매매가·보증금", "샘플 JSON", "데이터수집_및_API가이드_260714.md"],
        ["건축물대장", "주택유형·용도·면적 확인", "MVP", "건축물 특성", "수동 입력/업로드", "데이터수집_및_API가이드_260714.md"],
        ["OpenDART/사업자 상태", "법인 임대인 공개정보 보조", "확장", "공시·사업자 상태", "사용자 제출 서류", "데이터수집_및_API가이드_260714.md"],
        ["온비드", "공매 기반 대안 회수채널 분석", "확장", "공매 물건·입찰정보", "CSV/샘플 데이터", "데이터수집_및_API가이드_260714.md"],
        ["알림 채널", "D-90·증빙·위험경보 알림", "MVP 일부", "발송 결과", "화면 알림/이메일", "Backend_API_명세서_260714.md"],
    ], [2.5, 4.2, 1.5, 2.8, 2.8, 4.0], 6.4)
    para(doc, "API별 인증키 발급, 비용, 호출 제한, 요청 파라미터, 응답 JSON, 정규화 로직은 개발설계보고서에 반복하지 않고 데이터수집_및_API가이드_260714.md에서 관리한다.")
    ref_box(doc, "7장 상세 문서 참조", ["외부 데이터와 API 요청·응답: 데이터수집_및_API가이드_260714.md", "외부 API Adapter와 오류코드: Backend_API_명세서_260714.md", "Mock Dataset 구조: 데이터수집_및_API가이드_260714.md"])

    h1(doc, "8. 블록체인 상위 설계")
    para(doc, "블록체인은 계약·검증·상태변경의 공동 감사장부이자 임대차 생애주기 신뢰 레이어다. 단순 파일 해시 저장소로 축소하지 않되, 법적 강제집행 기능을 수행한다고 과장하지 않는다.")
    table(doc, "블록체인 기록 대상", ["상위 역할", "기록 사건", "연결 시스템"], [
        ["계약 버전 공증", "초안, 보완계약, 최종 계약, HUG 제출본", "Contract Service, Object Storage"],
        ["등기/API 조회 스냅샷", "조회시각, 결과 해시, 검증기관, 위험등급", "External API Gateway"],
        ["위험조건 해소 이력", "근저당 존재, 말소 예정, 말소 확인, 위험등급 변경", "Risk Engine, Verification Service"],
        ["검증기관별 서명·상태", "임대인 제출, 아이엔 검토, CODEF/API 조회, HUG 확인", "Verification Service"],
        ["계약 상태변경 감사로그", "ContractStatus 변경", "Timeline Service"],
        ["D-90 반환준비 기록", "요청, 응답, 보완, 미제출, 반환곤란", "Scheduler, Return Plan"],
        ["사고 후 HUG 인계", "사고의심, 이행청구, 채권관리 인계, 회수분석", "Incident, Recovery"],
    ], [4.0, 7.5, 5.0], 7.0)
    bullet(doc, ["온체인 금지: 계약서 원문, 등기부 원문, 주민등록번호, 사업자등록번호 원문, 상세주소, 예금·소득정보, 상담본문, 세금자료 원문.", "블록체인이 수행하지 못하는 것: 허위문서 진실성 자동 판정, 근저당 말소, 세금체납 조회, 원화 잔금 자동이체, 압류, 경매개시, HUG 대위변제 승인, 법원 배당순위 결정, ML 위험예측.", "Solidity 함수·이벤트·권한·테스트·배포 스크립트는 Blockchain_설계서_260714.md로 이관한다."])
    ref_box(doc, "8장 상세 문서 참조", ["스마트컨트랙트·상태머신: Blockchain_설계서_260714.md", "블록체인 Adapter API: Backend_API_명세서_260714.md"])

    doc.add_page_break()
    h1(doc, "9. 전체 시스템 아키텍처")
    code_lines(doc, "논리 아키텍처", [
        "사용자 -> Frontend -> Backend API Gateway -> Service Layer",
        "  ├ 계약·사용자 서비스  ├ 위험진단 서비스  ├ 증빙·검증 서비스",
        "  ├ D-90 및 알림 서비스 ├ 사고·채권관리 서비스 ├ ML 추론 서비스",
        "  ├ RAG 서비스 ├ 외부 API Adapter └ Blockchain Adapter",
        "-> Data Layer",
        "  ├ PostgreSQL ├ pgvector ├ Object Storage ├ Model Artifact └ Mock Dataset",
        "-> External Systems",
        "  ├ CODEF ├ 도로명주소 ├ 국토부 실거래가 ├ 건축물대장",
        "  ├ OpenDART ├ 사업자등록 상태 ├ 온비드 └ 알림 채널",
        "-> Polygon Amoy",
    ])
    table(doc, "모듈 책임 정의", ["모듈명", "주요 책임", "입력", "출력", "의존 모듈", "오류 시 처리", "MVP", "상세 문서"], MODULES, [2.1, 3.5, 2.5, 2.4, 2.3, 2.4, 1.2, 3.2], 5.4)
    h2(doc, "9.3 동기 요청")
    para(doc, "사용자 입력, 계약 조회, 위험진단 결과 확인, 증빙 제출, 관리자 대시보드 조회는 Frontend에서 Backend API Gateway로 동기 요청한다.")
    h2(doc, "9.4 비동기 처리")
    para(doc, "문서 추출, 외부 API 재시도, 알림 발송, 블록체인 트랜잭션 확인, 모델 배치 갱신은 작업 큐 또는 스케줄러로 비동기 처리한다.")
    h2(doc, "9.5 외부 API 흐름")
    para(doc, "Backend의 External API Gateway가 주소 정규화, CODEF 조회, 국토부·건축물대장 등 호출을 담당한다. 실패하면 재시도 후 MockFallback 또는 문서 업로드로 전환한다.")
    h2(doc, "9.6 ML 추론 흐름")
    para(doc, "Risk Engine 또는 Recovery Service가 ML Inference Service를 호출하고, 모델 결과와 SHAP 근거를 위험진단 또는 HUG 대시보드에 연결한다.")
    h2(doc, "9.7 블록체인 흐름")
    para(doc, "검증 완료, 계약 버전 확정, D-90 이벤트, 사고 인계 등 감사 대상 사건은 Blockchain Adapter를 통해 Polygon Amoy에 기록하고 tx hash를 DB에 저장한다.")
    h2(doc, "9.8 D-90 흐름")
    code_lines(doc, "D-90 스케줄러 흐름", ["PostgreSQL 계약 만료일 -> 백엔드 스케줄러 -> D-90 대상 계약 조회", "-> D-90 이벤트 생성 -> DB 저장 -> 임대인·임차인 알림", "-> 상태 변경 -> 상태변경 해시 또는 이벤트를 블록체인 기록"])
    h2(doc, "9.9 배포 구조")
    para(doc, "프론트엔드는 Vercel, 백엔드는 Render 또는 Railway, 데이터베이스와 스토리지는 Supabase 또는 로컬 PostgreSQL/S3 호환 스토리지를 우선 검토한다. Secret은 환경변수 기반으로 관리한다.")
    ref_box(doc, "9장 상세 문서 참조", ["내부 API·DB·ERD: Backend_API_명세서_260714.md", "화면·Wireframe·컴포넌트: Frontend_UIUX_명세서_260714.md", "외부 API Adapter: 데이터수집_및_API가이드_260714.md"])

    h1(doc, "10. 공통 데이터 계약")
    h2(doc, "10.1 공통 식별자")
    table(doc, "공통 식별자", ["식별자", "생성 시스템", "의미", "상세 정의 문서"], IDS, [3.2, 4.0, 5.5, 4.0], 7.0)
    h2(doc, "10.2 공통 상태")
    table(doc, "공통 상태", ["상태 그룹", "값", "상세 전이 정의"], [
        ["ContractStatus", "Draft, Diagnosed, EvidenceRequested, EvidenceSubmitted, Verified, ContractFinalized, Monitoring, D90Requested, ReturnPlanSubmitted, AtRisk, IncidentReported, TransferredToHUG, RecoveryInProgress, Closed", "Backend_API_명세서_260714.md, Blockchain_설계서_260714.md"],
        ["VerificationStatus", "Pending, Submitted, Reviewing, Verified, Rejected, Expired", "Backend_API_명세서_260714.md"],
        ["BlockchainStatus", "NotRequested, Pending, Confirmed, Failed", "Blockchain_설계서_260714.md"],
        ["APIResultStatus", "Success, Partial, Failed, MockFallback", "데이터수집_및_API가이드_260714.md, Backend_API_명세서_260714.md"],
        ["ModelResultStatus", "Success, RuleOnlyFallback, Failed, InsufficientData", "ML개발가이드_260714.md"],
    ], [3.2, 8.6, 4.7], 6.8)
    h2(doc, "10.3 공통 이벤트")
    bullet(doc, ["ContractCreated", "RegistrySnapshotCreated", "RiskAssessed", "EvidenceRequested", "EvidenceSubmitted", "VerificationCompleted", "ContractVersionFinalized", "D90Requested", "ReturnPlanSubmitted", "IncidentReported", "TransferredToHUG", "RecoveryPredictionCreated", "BlockchainConfirmed"])
    h2(doc, "10.4 API·DB·블록체인 참조 ID")
    para(doc, "API 응답, DB 레코드, 블록체인 이벤트는 동일한 업무 객체를 참조할 때 위 공통 식별자를 사용한다. 블록체인에는 원문 ID가 아닌 해시 또는 내부 참조값을 저장할 수 있으며 매핑 규칙은 Backend와 Blockchain 상세 문서에서 정의한다.")
    h2(doc, "10.5 시간·금액·주소 표준")
    bullet(doc, ["시간: KST 기준 ISO 8601", "금액: KRW 정수 원 단위", "주소: 도로명주소, 지번주소, 법정동코드, 건물관리번호, 동·호를 구분", "문서 해시: SHA-256 또는 컨트랙트 설계에서 지정한 bytes32 포맷"])

    h1(doc, "11. 상위 User Flow")
    table(doc, "사용자별 상위 흐름", ["사용자", "흐름", "상세 문서"], [
        ["임차인", "홈 -> 주소·계약 입력 -> API 조회 -> 위험진단 결과 -> 증빙요청 확인 -> D-90 -> 사고 대응", "Frontend_UIUX_명세서_260714.md"],
        ["임대인", "알림 -> 보완요청 확인 -> 증빙 제출 -> 검증상태 확인 -> 반환계획 제출", "Frontend_UIUX_명세서_260714.md"],
        ["아이엔", "상담 접수 -> RAG 근거 확인 -> 검증 검토 -> 전문가 이관", "ML개발가이드_260714.md, Frontend_UIUX_명세서_260714.md"],
        ["HUG", "대시보드 -> 사건 상세 -> 회수등급·처리기간 확인 -> 담당자 액션 -> 종결", "Backend_API_명세서_260714.md"],
        ["API 실패", "자동조회 -> 재시도 -> 업로드 또는 Mock -> 사용자 확인 -> 위험진단 계속", "데이터수집_및_API가이드_260714.md"],
        ["사고 후 인계", "IncidentReported -> TransferredToHUG -> RecoveryInProgress -> Closed", "Backend_API_명세서_260714.md"],
    ], [2.2, 9.6, 4.7], 7.0)

    h1(doc, "12. 상위 IA 및 UI/UX")
    h2(doc, "12.1 Sitemap")
    table(doc, "상위 IA", ["사용자", "화면 그룹"], [
        ["임차인", "홈, 계약 입력, API 조회, 위험결과, 계약 타임라인, D-90, 사고 대응"],
        ["임대인", "증빙 요청, 증빙 제출, 검증 진행, 반환계획 제출"],
        ["아이엔", "상담 RAG, 증빙 검토, 전문가 이관"],
        ["HUG", "대시보드, 사건 상세, 유사사건 비교, 법인 집단위험"],
        ["공통", "로그인, 알림, 권한 오류, 블록체인 이력"],
    ], [3.0, 13.5])
    h2(doc, "12.3 대표 화면")
    table(doc, "대표 화면 요약", ["대표 화면", "목적", "핵심 표시", "상세 문서"], [
        ["임차인 홈", "진단 시작과 계약 상태 확인", "진행 계약, 위험등급, 알림", "Frontend_UIUX_명세서_260714.md"],
        ["주소·계약정보 입력", "진단 기준 정보 수집", "주소 후보, 보증금, 기간", "Frontend_UIUX_명세서_260714.md"],
        ["API 조회 진행", "외부 API 상태 표시", "조회 단계, 재시도, 업로드 전환", "Frontend_UIUX_명세서_260714.md"],
        ["위험진단 결과", "등급·근거·다음 행동 제시", "위험요소, 증빙, 출처, 모델버전", "Frontend_UIUX_명세서_260714.md"],
        ["증빙·검증 진행", "보완요청과 검증 상태 관리", "요청사유, 제출, 검증결과", "Frontend_UIUX_명세서_260714.md"],
        ["계약 타임라인", "생애주기 상태 확인", "상태 이벤트, 블록체인 상태", "Frontend_UIUX_명세서_260714.md"],
        ["D-90 반환계획", "반환 준비 확인", "계획, 응답상태, 경보", "Frontend_UIUX_명세서_260714.md"],
        ["HUG 채권관리 대시보드", "사건 우선순위 관리", "조기경보, 회수등급, 처리기간", "Frontend_UIUX_명세서_260714.md"],
    ], [3.3, 4.2, 5.2, 3.8], 6.8)
    h2(doc, "12.4 UX 원칙")
    bullet(doc, ["점수보다 근거와 다음 행동을 우선한다.", "해결 가능한 위험과 해결 불가능한 위험을 분리한다.", "조회시각, 데이터 출처, 모델 버전, 블록체인 상태를 명확히 표시한다.", "API 실패 시 재시도, 업로드 대체, Mock 전환을 사용자에게 숨기지 않는다.", "블록체인 상태는 Pending/Confirmed/Failed로 표시하고 법적 효력으로 과장하지 않는다."])
    ref_box(doc, "12장 상세 문서 참조", ["화면·Wireframe·컴포넌트: Frontend_UIUX_명세서_260714.md", "내부 API 응답 상태: Backend_API_명세서_260714.md"])

    h1(doc, "13. 백엔드 및 데이터베이스 상위 설계")
    code_lines(doc, "백엔드 계층 구조", [
        "backend/",
        "├── api/             REST route",
        "├── core/            config, auth, logging",
        "├── schemas/         Pydantic DTO",
        "├── models/          SQLAlchemy model",
        "├── services/        business service",
        "├── repositories/    DB access",
        "├── integrations/    external API adapter",
        "├── ml/              inference loader",
        "├── rag/             retrieval service",
        "├── blockchain/      chain adapter",
        "├── scheduler/       D-90 and retry jobs",
        "├── notifications/   email/SMS/Kakao/FCM",
        "└── tests/           unit and integration tests",
    ])
    h2(doc, "13.2 핵심 서비스")
    para(doc, "핵심 서비스는 Contract, Property, Risk, Evidence, Verification, Timeline, Scheduler, Notification, Incident, Recovery, External API Gateway, Blockchain Adapter로 구성한다.")
    h2(doc, "13.3 핵심 엔터티")
    bullet(doc, ["User", "Property", "Contract", "RegistrySnapshot", "RiskAssessment", "EvidenceRequest", "Evidence", "Verification", "ContractVersion", "ReturnPlan", "Incident", "RecoveryPrediction", "TimelineEvent", "BlockchainTransaction"])
    h2(doc, "13.4 상위 ERD")
    code_lines(doc, "상위 엔터티 관계", [
        "User -> Contract -> Property -> RegistrySnapshot -> RiskAssessment",
        "Contract -> EvidenceRequest -> Evidence -> Verification",
        "Contract -> ContractVersion",
        "Contract -> ReturnPlan",
        "Contract -> Incident -> RecoveryPrediction",
        "Contract -> TimelineEvent -> BlockchainTransaction",
    ])
    h2(doc, "13.5 인증·권한")
    para(doc, "권한은 tenant, landlord, advisor, hug_admin, system_admin, verifier 역할을 기준으로 한다. 세부 권한 매트릭스와 인증 구현은 Backend_API_명세서_260714.md에서 정의한다.")
    h2(doc, "13.6 스케줄러")
    para(doc, "D-90 이벤트, API 재시도, 블록체인 확인, 알림 재발송은 APScheduler, Celery Beat, cron 또는 Cloud Scheduler 중 배포 환경에 맞춰 선택한다.")
    ref_box(doc, "13장 상세 문서 참조", ["내부 REST API·DB·ERD: Backend_API_명세서_260714.md", "외부 API Adapter 상세: 데이터수집_및_API가이드_260714.md"])

    doc.add_page_break()
    h1(doc, "14. 개발 및 통합 로드맵")
    table(doc, "6개 문서를 활용하는 개발 순서", ["순서", "작업", "참조/산출물"], [
        ["1", "개발설계보고서 확정", "본 문서"],
        ["2", "데이터수집·API 가이드 작성 및 API 키 신청", "데이터수집_및_API가이드_260714.md"],
        ["3", "Backend API·DB·ERD 명세 작성", "Backend_API_명세서_260714.md"],
        ["4", "ML 개발가이드 및 Colab Notebook 생성", "ML개발가이드_260714.md, .ipynb"],
        ["5", "Blockchain 설계 및 스마트컨트랙트 작성", "Blockchain_설계서_260714.md, Hardhat"],
        ["6", "Frontend UI/UX 명세 작성", "Frontend_UIUX_명세서_260714.md"],
        ["7", "Mock API 기반 Frontend 구현", "Mock Dataset, 화면 코드"],
        ["8", "Backend 구현", "FastAPI, PostgreSQL"],
        ["9", "ML Artifact 연결", "joblib/booster, model_versions"],
        ["10", "블록체인 연결", "Polygon Amoy, ethers.js"],
        ["11", "외부 API 연결", "CODEF, 공공 API"],
        ["12", "통합 테스트", "기능·API·권한·체인 테스트"],
        ["13", "시연 데이터·발표 준비", "시연 시나리오, 발표 메시지"],
    ], [1.4, 7.2, 7.9], 7.0)
    code_lines(doc, "문서 간 선후관계", [
        "개발설계보고서",
        "  -> 데이터수집/API 가이드 -> Backend API 명세 -> 구현",
        "  -> ML 개발가이드 -> Colab Notebook -> Model Artifact -> Backend 연결",
        "  -> Blockchain 설계서 -> Contract/ABI -> Backend 연결",
        "  -> Frontend UI/UX 명세서 -> 화면 구현 -> Backend/Mock 연결",
    ])
    table(doc, "단계별 완료 체크포인트", ["단계", "완료 조건", "다음 단계 진입 조건"], [
        ["문서 확정", "6개 문서의 책임 경계와 공통 식별자·상태 확정", "상세 명세서 작성 착수"],
        ["데이터/API", "필수 API 목록, 인증 필요 여부, Mock JSON 구조 확정", "Backend Adapter 설계"],
        ["Backend 명세", "핵심 엔터티, REST API 그룹, 오류코드, 권한 매트릭스 확정", "프론트·ML·체인 연동 계약 확정"],
        ["ML 명세", "모델 목적, 입력/출력, 평가기준, 저장 형식 확정", "Colab Notebook 생성"],
        ["Blockchain 명세", "기록 대상, 온체인 금지, 이벤트, role, 테스트넷 배포 기준 확정", "Hardhat 구현"],
        ["Frontend 명세", "대표 화면, 상태 UI, fallback UX, 접근성 기준 확정", "Mock API 기반 화면 구현"],
        ["통합", "계약 입력부터 위험진단, 검증, 체인 기록, D-90, HUG 대시보드까지 연결", "시연 리허설"],
        ["발표", "데이터 한계와 구현/확장 범위가 발표자료에 반영", "최종 제출"],
    ], [3.0, 7.0, 6.5], 7.0)

    h1(doc, "15. 해커톤 시연 시나리오")
    table(doc, "시연 시나리오", ["시나리오", "시연 내용", "강조 포인트"], [
        ["계약 전 위험 발견", "주소·계약 입력 후 등기·시세 기반 위험등급 표시", "점수보다 근거와 조치"],
        ["근저당 말소 검증", "근저당 발견 -> 보완요청 -> 말소 확인", "위험 해소 증명"],
        ["계약 버전 공증", "초안·보완·최종·HUG 제출본 해시 비교", "버전 불일치 방지"],
        ["D-90 조기경보", "스케줄러 이벤트와 반환계획 미응답 경보", "블록체인은 스케줄러가 아님"],
        ["사고 후 채권관리", "사고 인계 후 회수등급·처리기간 표시", "업무 우선순위 보조"],
        ["외부 API 실패 fallback", "CODEF 실패 -> PDF 업로드 또는 Mock JSON", "시연 안정성"],
    ], [3.5, 7.5, 5.5], 7.2)

    h1(doc, "16. 테스트 전략")
    table(doc, "상위 테스트 전략", ["분류", "검증 대상"], [
        ["기능", "계약 생성, 위험진단, 증빙 요청, 검증, 타임라인, D-90, 사고 접수"],
        ["통합", "Frontend-Backend-DB-ML-Blockchain-External API 연결"],
        ["API", "성공, Partial, Failed, MockFallback, timeout, 인증 실패"],
        ["모델", "입력 스키마, 결측, RuleOnlyFallback, InsufficientData, model_version"],
        ["블록체인", "상태 이벤트, 권한, Pending/Confirmed/Failed, 중복 기록"],
        ["권한", "tenant, landlord, advisor, hug_admin, admin 접근 제한"],
        ["개인정보", "원문 온체인 금지, 로그 마스킹, 파일 접근권한"],
        ["시연 안정성", "Mock 모드, 사전 샘플, 재시도 버튼, 오류 메시지"],
    ], [3.0, 13.5])

    h1(doc, "17. 보안·개인정보·법률")
    bullet(doc, ["최소수집: 진단과 검증에 필요한 정보만 수집한다.", "원문 저장: 계약서·등기부·증빙 원문은 Object Storage에 권한 기반으로 저장한다.", "온체인 금지: 개인정보와 문서 원문은 블록체인에 저장하지 않는다.", "비밀키: 지갑 개인키와 API Secret은 환경변수 또는 Secret Manager로 관리한다.", "API 키: 프론트엔드에 노출하지 않고 Backend Adapter에서만 사용한다.", "법률정보와 자문 구분: 서비스 설명은 정보 제공이며 법률자문이 아니다.", "금융·법원 기능 한계: 원화 이체, 압류, 경매개시, 대위변제 승인, 배당순위 결정은 수행하지 않는다."])

    h1(doc, "18. 리스크와 대응")
    table(doc, "리스크와 대응", ["리스크", "영향", "대응"], [
        ["정상계약 데이터 부재", "실제 사고확률 모델 불가", "위험유사도·고위험 패턴·상대위험도로 표현"],
        ["비식별 법인키 부재", "동일 법인 다물건 그래프 불가", "법인/지역/주택유형 집단 분석으로 제한"],
        ["API 비용·승인", "실 API 연결 지연", "Mock API와 샘플 JSON 병행"],
        ["API 장애", "진단 지연", "재시도, 업로드 대체, MockFallback"],
        ["법인 공시 누락", "법인 위험 판단 제한", "OpenDART 누락 가능성 표시"],
        ["경매 데이터 부족", "회수모델 정확도 제한", "발제사에 회차별 데이터 요청"],
        ["임대인 참여", "증빙 제출 지연", "요청 사유와 계약 진행 영향 명확화"],
        ["모델 일반화", "표본 편향", "시간 기반 검증과 SHAP 설명"],
        ["블록체인 과대적용", "법적 효력 오해", "감사장부 역할로 명확히 제한"],
    ], [3.6, 6.0, 6.9], 7.2)

    h1(doc, "19. 평가기준 대응")
    table(doc, "평가기준 대응", ["기준", "대응"], [
        ["독창성", "위험 점수 앱이 아니라 위험 해소 검증과 사고 전후 연결을 구현"],
        ["데이터 활용", "등기, 실거래가, 건축물, 상담, 사고·배당 데이터를 역할별로 연결"],
        ["기술성", "FastAPI, RAG, ML, Polygon Amoy, CODEF/API Gateway 통합"],
        ["실현 가능성", "MVP와 Mock, PoC, 확장 범위를 분리"],
        ["확장성", "비식별 법인키, 기관 연계, 에스크로, 채권관리 고도화로 확장 가능"],
        ["공공성", "임차인 보호와 HUG 업무 우선순위 보조"],
    ], [3.0, 13.5])

    h1(doc, "20. 최종 결론")
    para(doc, "HUG 안심전세 체인은 계약 전 위험진단, 위험조건 보완, 외부 API·증빙 검증, 계약 버전과 검증 이력 공증, D-90 조기개입, 사고 후 HUG 채권관리 인계를 하나의 생애주기로 연결한다. 본 개발설계보고서는 상세 구현을 모두 담는 문서가 아니라 5개 Markdown 상세 명세서가 충돌 없이 작성되고 코드 생성으로 이어지도록 프로젝트 전체 기준과 공통 계약을 정의하는 최상위 문서다.")

    doc.add_page_break()
    h1(doc, "부록 A. 상세 문서 목록")
    table(doc, "상세 문서 목록", ["파일명", "상태"], [[r[0], "향후 작성 예정" if r[7] == "Markdown" else "본 문서의 이전본"] for r in DOCS], [8.0, 8.5])
    h1(doc, "부록 B. 공통 식별자")
    table(doc, "공통 식별자 재수록", ["식별자", "생성 시스템", "상세 정의 문서"], [[r[0], r[1], r[3]] for r in IDS], [4.8, 5.6, 6.1], 7.0)
    h1(doc, "부록 C. 공통 상태")
    bullet(doc, ["ContractStatus: Draft, Diagnosed, EvidenceRequested, EvidenceSubmitted, Verified, ContractFinalized, Monitoring, D90Requested, ReturnPlanSubmitted, AtRisk, IncidentReported, TransferredToHUG, RecoveryInProgress, Closed", "VerificationStatus: Pending, Submitted, Reviewing, Verified, Rejected, Expired", "BlockchainStatus: NotRequested, Pending, Confirmed, Failed", "APIResultStatus: Success, Partial, Failed, MockFallback", "ModelResultStatus: Success, RuleOnlyFallback, Failed, InsufficientData"])
    h1(doc, "부록 D. 모듈 책임표")
    table(doc, "모듈 책임표 재수록", ["모듈명", "주요 책임", "MVP", "상세 문서"], [[m[0], m[1], m[6], m[7]] for m in MODULES], [3.3, 7.2, 1.5, 4.5], 6.8)
    h1(doc, "부록 E. MVP·PoC·확장 구분")
    table(doc, "구분표", ["구분", "대상"], [["MVP", "계약 입력, API/업로드 조회, 위험진단, 증빙·검증, 계약 버전, D-90 데모, 관리자 Mock"], ["PoC", "HUG 데이터, CODEF 실계정, 회수모델, 알림 채널"], ["확장", "비식별 법인키 그래프, 은행 에스크로, 기관 자동연계"], ["미구현", "법적 강제집행, 원화 자동이체, 법원 배당순위 결정"]], [3.0, 13.5])
    h1(doc, "부록 F. 발제사 확인 질문")
    bullet(doc, ["비식별 법인키 제공 가능 여부", "정상계약 데이터 제공 여부와 기준시점", "공통 계약·물건·사고키 존재 여부", "경매 회차별 데이터와 배당 데이터 범위", "HUG 기존 등기 모니터링 또는 알림 업무 존재 여부", "D-90 반환계획 확인 제도화 가능성", "채권관리 우선순위 기준과 현재 업무량·처리비용", "블록체인 감사장부에 대한 실무 수요", "외부 API 및 Mock 사용 허용 범위", "CODEF 실제 계정, 과금, 호출 제한, 테스트 가능 여부"])

    doc.core_properties.title = "HUG × 아이엔 안심주거 생태계 개발설계보고서 수정보완"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
