from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from pathlib import Path
import json
import zipfile


BASE = Path("/Users/choijihee/Library/CloudStorage/GoogleDrive-heeppiness@pukyong.ac.kr/내 드라이브/DIVE2026")
SOURCE = BASE / "개발설계보고서_260712.docx"
OUT = BASE / "개발설계보고서_260714.docx"
COMMENTS_JSON = Path("/private/tmp/dive2026_comments.json")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, size=8.5, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(text))
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "목차는 Word에서 문서 열기 시 필드 업데이트로 갱신됩니다."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def style_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.1)
    sec.right_margin = Cm(2.1)
    sec.header_distance = Cm(1.2)
    sec.footer_distance = Cm(1.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Title", 22, "0B2545", 0, 10),
        ("Subtitle", 11, "334155", 0, 10),
        ("Heading 1", 15, "0B2545", 14, 7),
        ("Heading 2", 12, "1F4D78", 10, 5),
        ("Heading 3", 10.5, "334155", 8, 3),
    ]:
        st = styles[name]
        st.font.name = "맑은 고딕"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = name.startswith("Heading") or name == "Title"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = name.startswith("Heading")

    if "CaptionText" not in styles:
        st = styles.add_style("CaptionText", WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = "맑은 고딕"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        st.font.size = Pt(8.5)
        st.font.color.rgb = RGBColor.from_string("475569")
        st.paragraph_format.space_before = Pt(4)
        st.paragraph_format.space_after = Pt(2)

    if "CodeBlock" not in styles:
        st = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = "Courier New"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        st.font.size = Pt(8.5)
        st.paragraph_format.space_before = Pt(2)
        st.paragraph_format.space_after = Pt(2)

    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "HUG × 아이엔 안심주거 생태계 개발설계보고서"
        header.style = doc.styles["Normal"]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer = section.footer.paragraphs[0]
        footer.text = "DIVE 2026 | 2026.07.14 | "
        add_page_number(footer)


table_no = 0
fig_no = 0


def caption(doc, text, kind="표"):
    global table_no, fig_no
    if kind == "표":
        table_no += 1
        label = f"표 {table_no}. {text}"
    else:
        fig_no += 1
        label = f"그림 {fig_no}. {text}"
    p = doc.add_paragraph(label, style="CaptionText")
    return p


def add_table(doc, headers, rows, widths=None, font_size=8.0):
    if headers and headers[0].startswith("표:"):
        caption_title = headers[0].replace("표:", "", 1).strip()
    elif headers:
        caption_title = f"{headers[0]} 기준 정리"
    else:
        caption_title = "정리"
    caption(doc, caption_title)
    if headers and headers[0].startswith("표:"):
        headers = headers[1:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, bold=True, size=font_size, color="0B2545")
        set_cell_shading(hdr.cells[i], "E8EEF5")
        if widths:
            set_cell_width(hdr.cells[i], widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=font_size)
            if widths:
                set_cell_width(cells[i], widths[i])
    return table


def add_kv_table(doc, title, pairs):
    caption(doc, title)
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    for k, v in pairs:
        row = table.add_row()
        set_cell_text(row.cells[0], k, bold=True, size=8.5, color="0B2545")
        set_cell_shading(row.cells[0], "F2F4F7")
        set_cell_text(row.cells[1], v, size=8.5)
        set_cell_width(row.cells[0], 4.2)
        set_cell_width(row.cells[1], 12.2)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def para(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def add_flow(doc, title, lines):
    caption(doc, title, kind="그림")
    for line in lines:
        p = doc.add_paragraph(style="CodeBlock")
        p.add_run(line)


def comments_summary():
    if not COMMENTS_JSON.exists():
        return "원본 댓글 파일을 찾지 못함"
    data = json.loads(COMMENTS_JSON.read_text())
    comments = data.get("comments", [])
    return "; ".join([f"{c.get('id')}: {c.get('text','').replace(chr(10), ' / ')}" for c in comments])


api_rows = [
    ["CODEF 부동산 등기부등본", "CODEF", "주소 기반 등기 조회, 갑구·을구 위험항목 구조화", "도로명/지번주소, 동·호, 사용자 동의", "소유자, 압류, 가압류, 경매, 근저당, 채권최고액, 조회일시", "필요", "확인 필요", "MVP 핵심", "PDF 업로드, 샘플 JSON", "원문·상세주소 온체인 금지, 계정·과금 확인"],
    ["도로명주소 API", "행정안전부/도로명주소", "주소 정규화 및 기준키 생성", "주소 문자열", "도로명주소, 지번주소, 법정동코드, 건물관리번호", "확인 필요", "낮음/확인 필요", "MVP 핵심", "수동 주소 입력·후보 선택", "주소 후보 다중 선택 UI 필요"],
    ["아파트 매매 실거래가", "국토교통부", "주변 매매가와 전세가율 보조", "법정동, 기간, 단지", "거래금액, 면적, 층, 계약월", "공공데이터 키", "대체로 무료/확인 필요", "MVP", "사전 저장 샘플", "시장가격 확정값 아님"],
    ["아파트 전월세 실거래가", "국토교통부", "주변 보증금 수준 비교", "법정동, 기간, 단지", "보증금, 월세, 면적, 층", "공공데이터 키", "대체로 무료/확인 필요", "MVP", "샘플 JSON", "동일 조건 비교 한계 표시"],
    ["연립·다세대 매매/전월세", "국토교통부", "빌라 유형 비교와 이상계약 탐지", "법정동, 기간", "거래금액, 보증금, 면적", "공공데이터 키", "확인 필요", "MVP", "지역 평균 대체", "다가구/다세대 구분 주의"],
    ["단독·다가구 매매/전월세", "국토교통부", "다가구 선순위보증금 판단 보조", "법정동, 기간", "거래/임대차 사례", "공공데이터 키", "확인 필요", "PoC", "수동 입력", "선순위보증금 전체 파악 불가"],
    ["오피스텔 매매·전월세", "국토교통부", "오피스텔 전세가율 비교", "법정동, 기간", "거래금액, 보증금", "공공데이터 키", "확인 필요", "MVP", "샘플 JSON", "주거용/업무용 구분 확인"],
    ["건축물대장", "국토교통부/공공데이터포털", "용도, 면적, 사용승인, 위반건축물 항목 확인", "주소, 건물관리번호", "표제부, 전유부, 층별개요, 용도", "공공데이터 키", "확인 필요", "MVP", "사용자 업로드/수동 입력", "제공 항목 범위 확인 필요"],
    ["공동주택 공시가격", "한국부동산원/국토부", "가격 보조지표", "주소, 단지, 동·호", "공시가격", "확인 필요", "확인 필요", "PoC", "실거래가 우선", "실시간 시장가격 아님"],
    ["개별주택가격·개별공시지가", "지자체/국토부", "단독·다가구 보조가격", "주소, 지번", "주택가격, 토지가격", "확인 필요", "확인 필요", "확장", "유사거래 비교", "보조지표로만 사용"],
    ["OpenDART", "금융감독원", "법인 임대인 공개정보 탐색", "법인명/사업자번호/고유번호", "기업개황, 재무, 감사보고서, 주요공시", "API 키", "무료 한도/확인 필요", "확장", "사용자 제출 서류", "비상장 소형 법인은 누락 가능"],
    ["사업자등록 상태조회", "국세청", "계속·휴업·폐업 확인", "사업자등록번호, 동의", "사업자 상태", "인증/동의 확인", "확인 필요", "PoC", "수동 증빙", "사업자번호 원문 보호"],
    ["온비드 Open API", "한국자산관리공사", "공매 물건·입찰정보와 대안 회수채널 분석", "지역, 물건유형, 기간", "공매 물건, 입찰, 낙찰 정보", "API 키 확인", "확인 필요", "확장", "CSV/샘플 데이터", "법원 경매와 모집단 분리"],
    ["알림 채널", "이메일/SMS/카카오/FCM", "D-90, 증빙요청, 위험경보", "수신자, 템플릿, 동의", "발송 결과, 실패 사유", "채널별 필요", "유료 가능", "MVP 일부", "이메일 우선, 화면 알림", "수신동의와 로그 관리"],
    ["Polygon Amoy", "Polygon/Alchemy 등 RPC", "계약·검증·상태변경 감사장부", "해시, 상태, 서명자", "트랜잭션, 이벤트, 블록번호", "RPC/지갑 필요", "테스트넷 무료성 확인", "MVP", "DB 감사로그", "개인정보·원문 온체인 금지"],
]


screen_rows = [
    ["TEN-01", "임차인 홈", "임차인", "계약 진단 시작과 현재 계약 상태 확인", "로그인 후", "계약 기본정보", "GET /contracts", "진행중 계약, 위험등급, D-90 여부", "상태 카드, CTA, 알림", "계약 진단 시작", "Default/Loading/Empty/API failure", "TEN-02", "tenant", "조회 실패", "하단 CTA 고정"],
    ["TEN-02", "주소·계약정보 입력", "임차인", "주소와 보증금 등 진단 입력", "TEN-01", "주소, 보증금, 기간, 임대인 유형", "POST /contracts, 주소 API", "주소 후보, 입력 검증", "주소검색, 금액 입력, 단계표시", "다음", "Validation error/Partial data", "TEN-03", "tenant", "주소 후보 없음", "주소 검색 모달 전체화면"],
    ["TEN-03", "API 조회 진행", "임차인", "외부 API 조회 상태 표시", "TEN-02", "계약 ID", "POST /registry/search, /market", "조회 단계별 상태", "진행바, 재시도, 업로드 전환", "조회/재시도", "Loading/API failure/Partial data", "TEN-04 또는 TEN-17", "tenant", "CODEF 실패", "단계별 세로 타임라인"],
    ["TEN-04", "위험진단 결과", "임차인", "위험등급과 근거, 해결 가능성 제시", "TEN-03", "contract_id", "GET /risk", "등급, 위험요소, 근거, 증빙, 모델버전, 출처", "등급 배지, 근거 카드, CTA", "증빙 요청", "Success/Model inference failure/Blockchain pending", "TEN-05", "tenant", "모델 실패 시 규칙 결과만 표시", "점수보다 조치 CTA 우선"],
    ["TEN-05", "위험요소 상세", "임차인", "근저당·압류 등 요소별 해석", "TEN-04", "risk_id", "GET /risk/{id}", "해결가능/불가능, 증빙, 계약 권고", "아코디언, 출처표, 상담 CTA", "전문가 상담", "Default/Partial data", "TEN-06", "tenant", "근거 데이터 누락", "긴 설명 접기"],
    ["LAND-01", "증빙 요청", "임대인", "보완조건과 제출 증빙 확인", "알림 링크", "request_id", "GET /evidence-requests", "요청사유, 기한, 파일요건", "업로드, 체크리스트", "증빙 제출", "Permission denied/Validation error", "LAND-02", "landlord", "권한 없음", "파일선택 단순화"],
    ["LAND-02", "검증 진행", "임대인/아이엔", "증빙 검토와 API 대조", "LAND-01", "evidence_id", "POST /verify", "검증상태, 보완요청", "상태 타임라인, 코멘트", "보완 제출", "Loading/Success/Rejected", "TEN-07", "landlord,advisor", "문서 추출 실패", "상태 중심"],
    ["TEN-07", "계약 버전 비교", "임차인/임대인", "초안·보완·최종 계약 차이 확인", "검증 완료", "contract_id", "GET /versions", "버전, 해시, 차이, 제출처", "버전 탭, diff 요약", "최종 확정", "Blockchain pending/confirmed", "TEN-08", "tenant,landlord", "버전 불일치", "차이 요약 먼저"],
    ["TEN-08", "블록체인 이력", "공통", "계약·검증 이력 공증 확인", "계약 상세", "contract_id", "GET /blockchain/events", "tx hash, block, status, verifier", "타임라인, 탐색기 링크", "탐색기 열기", "Blockchain pending/confirmed/API failure", "TEN-09", "authorized", "RPC 오류", "해시는 축약표시"],
    ["TEN-09", "계약 타임라인", "공통", "계약 생애주기 상태 확인", "계약 상세", "contract_id", "GET /timeline", "상태, 증빙, 알림, D-90", "상태머신, 필터", "다음 조치", "Default/Empty", "TEN-10", "authorized", "상태 누락", "세로형"],
    ["TEN-10", "D-90 반환계획", "임차인/임대인", "반환계획 요청·응답·증빙 관리", "D-90 알림", "contract_id", "GET/POST /return-plans", "계획, 방법, 증빙, 위험경보", "폼, 증빙 업로드, 경고", "제출/보완요청", "Validation error/AtRisk", "TEN-11", "tenant,landlord", "미응답", "마감일 강조"],
    ["TEN-11", "사고 대응", "임차인", "미반환 의심 또는 사고 접수", "타임라인", "contract_id, 사유", "POST /incidents", "접수상태, 필요서류", "사고접수 폼, 안내", "HUG 인계", "Success/Permission denied", "HUG-01", "tenant", "중복 접수", "단계형 폼"],
    ["HUG-01", "HUG 대시보드", "HUG 담당자", "조기경보와 사고 우선순위 관리", "관리자 로그인", "필터", "GET /admin/dashboard", "조기경보, 사고, 집중관리, 회수율, 처리기간, 채권규모", "KPI, 필터, 사건표", "사건 열기", "Loading/Partial data", "HUG-02", "hug_admin", "권한 없음", "핵심 KPI 2열"],
    ["HUG-02", "사건 상세", "HUG 담당자", "사건 근거와 회수 예측 확인", "HUG-01", "incident_id", "GET /incidents/{id}", "예상 회수율, 처리기간, 유사사건, 액션", "근거 패널, 우선순위", "담당자 배정", "Model inference failure/Partial data", "HUG-03", "hug_admin", "예측 실패", "액션 버튼 고정"],
    ["HUG-03", "유사사건 비교", "HUG 담당자", "사례군 비교와 조건 차이 분석", "HUG-02", "incident_id", "GET /similar-cases", "유사사건, 차이, 결과", "비교표, 필터", "분석 저장", "Empty/Partial data", "HUG-02", "hug_admin", "유사사례 없음", "표는 가로 스크롤"],
    ["HUG-04", "법인 위험분석", "HUG/아이엔", "비식별키 없는 집단 위험 분석", "대시보드", "지역, 유형, 금액구간", "GET /landlord-groups", "사고사례군 내 상대 위험도", "필터, 분포 차트", "보고서 생성", "Default/Empty", "HUG-02", "advisor,hug_admin", "표본 부족", "확률 표현 금지"],
    ["TEN-17", "API 실패·업로드 대체", "임차인", "자동조회 실패 시 PDF 업로드와 샘플 전환", "TEN-03", "PDF 또는 샘플 선택", "POST /documents/extract", "추출항목, 사용자 확인", "업로드, 추출확인표", "위험진단 계속", "API failure/Validation error/Success", "TEN-04", "tenant", "추출 실패", "사진 촬영 가이드"],
]


db_rows = [
    ["users", "id, role, name_hash, phone_hash, consent_flags", "역할·동의·비식별 사용자 관리"],
    ["properties", "id, address_hash, road_addr_key, building_mgmt_no, housing_type", "물건 기준키와 주소 정규화 결과"],
    ["contracts", "id, property_id, tenant_id, landlord_type, deposit_amount, start_date, end_date, status", "계약 생애주기 중심 엔티티"],
    ["registry_snapshots", "id, contract_id, source, queried_at, result_hash, mortgage_amount, seizure_flags", "등기 조회 시점 기준 스냅샷"],
    ["risk_assessments", "id, contract_id, grade, score_components_json, model_version, source_status", "위험등급과 근거"],
    ["evidence_requests", "id, contract_id, evidence_type, reason, due_date, status", "보완조건 요청"],
    ["evidences", "id, request_id, object_uri, document_hash, extracted_json, submitter_id", "증빙 원본은 오브젝트 스토리지, 해시는 DB/체인"],
    ["verifications", "id, evidence_id, verifier_role, result, reason, verified_at", "기관·담당자 검증 결과"],
    ["contract_versions", "id, contract_id, version_no, doc_hash, status, submitted_to_hug", "계약 버전 공증"],
    ["return_plans", "id, contract_id, requested_at, response_status, plan_type, risk_flag", "D-90 반환계획"],
    ["incidents", "id, contract_id, incident_type, claim_status, subrogation_amount, transferred_at", "사고 접수와 HUG 인계"],
    ["recovery_predictions", "id, incident_id, expected_recovery_grade, expected_days, model_version, shap_json", "회수등급·처리기간 예측"],
    ["timeline_events", "id, contract_id, event_type, payload_json, actor_role, created_at", "상태 변경 이력"],
    ["blockchain_transactions", "id, ref_type, ref_id, tx_hash, chain_id, status, block_no", "온체인 기록 추적"],
    ["api_call_logs", "id, provider, endpoint, status, retry_count, fallback_used, created_at", "외부 API 장애·비용 추적"],
    ["model_versions", "id, model_name, artifact_uri, metrics_json, train_data_version, created_at", "모델 버전관리"],
]


def build():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("HUG × 아이엔 안심주거 생태계\n개발기획 및 상세설계보고서")
    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("계약 전 위험진단부터 계약 이행 검증, 만료 전 조기개입, 사고 후 채권관리까지 연결하는 전세 계약 생애주기 플랫폼")
    add_kv_table(doc, "표지 정보", [
        ("프로젝트 가칭", "HUG 안심전세 체인"),
        ("작성 기준일", "2026년 7월 14일(KST)"),
        ("발제사", "HUG × 아이엔"),
        ("문서 성격", "해커톤 프로토타입 개발, UI/UX 생성, 백엔드·ML·RAG·블록체인 구현의 공통 기준 문서"),
        ("핵심 메시지", "위험을 점수화하는 앱이 아니라, 계약 위험이 실제로 해소됐는지 증명하고 사고 전후의 대응을 끊김 없이 연결하는 주거안전 인프라"),
    ])
    doc.add_page_break()

    h1(doc, "문서 변경이력")
    add_table(doc, ["버전", "작성일", "변경사항"], [[
        "1.0", "2026.07.14",
        "외부 API 연계 구조 구체화; CODEF 주소 기반 등기조회 반영; Google Colab ML 환경 반영; 비식별 법인키 부재 가정 반영; 블록체인 계약 상태 및 검증이력 역할 확대; UI/UX 화면 정의 보강; MVP·PoC·확장 기능 재분류"
    ]], widths=[2.0, 3.0, 11.5], font_size=8.5)
    add_kv_table(doc, "원본 검토 및 댓글 반영 요약", [
        ("사용 원본", str(SOURCE)),
        ("추출 댓글", comments_summary()),
        ("최종 처리", "댓글의 취지는 본문에 통합하고 최종 DOCX에는 Word 댓글을 남기지 않는다."),
    ])
    doc.add_page_break()

    h1(doc, "목차")
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    h1(doc, "0. 문서 개요")
    add_kv_table(doc, "문서 개요", [
        ("프로젝트명", "HUG × 아이엔 안심주거 생태계 / HUG 안심전세 체인"),
        ("대응 소주제", "전세 계약 위험진단, 계약 이행 검증, 만료 전 조기개입, 사고 후 채권관리 지원"),
        ("문서 목적", "해커톤 MVP 개발과 발표·시연·UI/UX 생성에 필요한 공통 기준을 정의한다."),
        ("개발 원칙", "실제 구현 가능성, 데이터 근거, 개인정보 최소화, 온체인 원문 저장 금지, MVP와 확장 기능의 명확한 분리"),
        ("표현 제한", "정상계약 데이터와 사고계약 데이터가 동일 기준시점으로 연결되지 않으면 실제 사고확률이라고 표현하지 않는다."),
    ])
    add_table(doc, ["구분", "포함 범위", "제외 또는 유보 범위"], [
        ["MVP", "주소 정규화, CODEF 등기조회 또는 업로드 대체, 위험등급, 보완요청, 계약버전·검증해시, D-90 데모, HUG 대시보드 Mock", "은행 실이체, 법적 강제집행, 실제 HUG 내부망 연동"],
        ["Mock·데모", "Mock API, 샘플 JSON, 정상/근저당/압류 등기부 샘플, 가상 에스크로 상태전이", "데모 결과를 실제 예측성능으로 주장하지 않음"],
        ["기관 연계형 PoC", "CODEF 실계정, HUG 데이터 연계, 알림 채널, 업무 우선순위 검증", "운영 승인은 별도"],
        ["장기 확장", "비식별 법인키 기반 그래프 분석, 은행 에스크로, 법원·세금·금융기관 연계", "해커톤 단독 구현 범위 아님"],
    ], widths=[2.3, 7.2, 7.0])

    h1(doc, "1. 프로젝트 배경과 문제 정의")
    h2(doc, "1.1 발제 요구사항")
    para(doc, "발제의 핵심은 전세사기와 보증금 미반환 위험을 계약 전후 전 과정에서 조기에 발견하고, HUG와 상담·검증기관의 업무가 단절되지 않도록 데이터를 연결하는 것이다. 본 문서는 단순 위험 점수 앱이 아니라 위험 해소 조건의 제안, 검증, 이력 공증, 사고 후 채권관리 인계까지 포함하는 플랫폼 기준을 제시한다.")
    h2(doc, "1.2 현재 임대차 보호 구조의 단절")
    add_bullets(doc, [
        "계약 전 위험 진단은 등기·시세·건축물 정보가 분산되어 임차인이 직접 해석하기 어렵다.",
        "근저당 말소, 압류 해소, 특약 반영 등 보완조건이 실제로 이행됐는지 증명하는 공통 장부가 없다.",
        "계약 만료 전 반환계획 확인이 늦어 사고 직전까지 위험 신호가 누락된다.",
        "사고 후 HUG 채권관리 단계에서 계약 전 위험정보와 검증 이력이 충분히 이어지지 않는다.",
    ])
    h2(doc, "1.3 해결하려는 핵심 문제")
    add_flow(doc, "전세 계약 생애주기 문제 흐름", [
        "계약 전 위험진단 -> 위험조건 보완요청 -> 외부 API 및 증빙 검증 -> 계약서 및 검증결과 기록",
        "-> 계약기간 중 상태변화 추적 -> D-90 반환계획 확인 -> 이상징후 조기경보",
        "-> 보증금 미반환 사고 -> HUG 채권관리 우선순위 및 회수 가능성 분석",
    ])
    h2(doc, "1.4 프로젝트 목표")
    add_numbered(doc, [
        "위험 발견: 등기, 시세, 건축물, 임대인 유형, 상담사례 기반 위험요소를 구조화한다.",
        "위험 해소: 보완조건과 필요한 증빙을 생성한다.",
        "위험 검증: CODEF/API 조회 또는 업로드 문서 추출을 통해 해소 여부를 검증한다.",
        "이력 관리: 계약 버전, 조회결과, 검증상태, D-90, 사고 인계를 블록체인 감사장부에 기록한다.",
        "사고 후 지원: 예상 회수등급 또는 예상 처리기간으로 HUG 사건 우선순위를 보조한다.",
    ])
    h2(doc, "1.5 서비스 한 문장 정의")
    para(doc, "HUG 안심전세 체인은 전세 계약 위험을 발견하는 데서 멈추지 않고, 위험이 해소됐는지 검증하고, 계약 만료 전 반환 준비와 사고 후 채권관리까지 연결하는 전세 계약 생애주기 플랫폼이다.")
    h2(doc, "1.6 핵심 가치 제안")
    add_table(doc, ["가치", "설명", "구현 기준"], [
        ["증명 가능성", "계약서와 검증결과의 버전 일치를 확인", "문서 해시, 검증자, 상태변경 이벤트 기록"],
        ["조기개입", "D-90 반환계획 확인으로 사고 전 위험 신호 포착", "백엔드 스케줄러와 알림"],
        ["업무 연결", "상담·검증·HUG 채권관리 데이터 연결", "계약 ID와 사고 ID 중심 DB"],
        ["과장 방지", "데이터 한계를 명시", "실제 사고확률 표현 금지"],
    ], widths=[2.5, 6.5, 7.5])

    h1(doc, "2. 사용자 및 이해관계자 정의")
    add_table(doc, ["사용자", "Pain Point", "기대효과", "주요 화면"], [
        ["임차인", "등기·시세·계약 특약을 해석하기 어렵고 사고 전 대응 시점이 늦다.", "위험요소, 보완조건, 증빙 필요사항, 계약 진행 권고를 한 화면에서 확인", "임차인 홈, 위험진단 결과, D-90, 사고 대응"],
        ["임대인", "정상 계약임을 설명할 증빙 제출 경로가 분산되어 있다.", "보완요청과 증빙 제출 상태를 구조화", "증빙 요청, 검증 진행"],
        ["HUG 채권관리 담당자", "사고 접수 후 우선순위 판단과 유사사건 비교에 시간이 든다.", "예상 회수등급·처리기간·유사사건 기반 우선순위 보조", "HUG 대시보드, 사건 상세"],
        ["아이엔 상담 담당자", "상담 유형 분류와 전문가 이관 판단이 반복된다.", "분쟁유형·진행단계·이관 필요도 분류", "상담 RAG, 검증 검토"],
        ["외부 API 사업자", "서비스 내 조회 결과의 역할과 보관 범위가 명확해야 한다.", "조회 시점·결과 해시·기관명만 공증", "API Gateway, 로그"],
    ], widths=[3.0, 5.2, 5.2, 3.1])

    h1(doc, "3. 서비스 범위")
    add_table(doc, ["범위", "정의", "구현 수준"], [
        ["전체 생애주기", "계약 전, 계약 이행, 만료 전, 사고 후 채권관리까지 연결", "최종 목표"],
        ["MVP", "주소 기반 등기조회, 위험등급, 보완요청, 검증이력, D-90 데모, HUG Mock 대시보드", "해커톤 구현"],
        ["Mock·데모", "외부 API 장애 대응과 발표 안정성 확보용 샘플", "사전 저장 JSON"],
        ["기관 연계형 PoC", "HUG/아이엔 데이터와 CODEF 실계정 연계", "발제사 확인 후"],
        ["장기 정책 확장", "법인 다물건 그래프, 에스크로, 제도형 D-90", "정책·법률 검토 필요"],
        ["구현하지 않는 기능", "근저당 말소, 압류, 세금체납 조회, 대위변제 승인, 법원 배당순위 결정", "불가능 또는 외부기관 권한"],
    ], widths=[3.0, 8.5, 5.0])

    h1(doc, "4. 핵심 서비스 기능")
    for title, body in [
        ("4.1 계약 전 진단", "주소, 보증금, 임대인 유형, 계약기간, 주택유형을 입력받아 등기·시세·건축물 정보와 결합한다."),
        ("4.2 등기부 자동조회", "CODEF 부동산 등기부등본 API를 주소 기반 조회의 기본 경로로 설계한다. 자동조회 실패 시 PDF 업로드로 전환한다."),
        ("4.3 외부 시세·건축물 정보 결합", "국토부 실거래가와 건축물대장으로 전세가율, 용도, 다가구·다세대 구분, 사용승인일을 보조 판단한다."),
        ("4.4 위험등급 및 위험사유", "A/B/C/D 등급과 위험요소별 근거를 분리하여 제시한다. 점수만 표시하지 않는다."),
        ("4.5 보완조건 생성", "근저당 말소 확인, 계약서 특약 반영, 추가 증빙 제출 등 해결 가능한 위험을 요청한다."),
        ("4.6 증빙 요청·검증", "임대인 제출, 아이엔 검토, API 조회 결과, HUG 확인 상태를 검증 이력으로 저장한다."),
        ("4.7 계약 버전 관리", "초안, 보완계약, 최종 계약, HUG 제출 계약서의 해시와 버전 일치를 확인한다."),
        ("4.8 계약기간 모니터링", "1회 조회와 별개로 주기적 재조회, 비용, 동의, 변경 비교, 알림을 갖춘 확장 기능으로 설계한다."),
        ("4.9 D-90 반환계획", "스마트컨트랙트가 아니라 백엔드 스케줄러가 만료일을 조회해 이벤트를 생성한다."),
        ("4.10 사고 대응", "미반환 의심 접수, 필요서류 안내, HUG 인계 상태를 타임라인에 기록한다."),
        ("4.11 HUG 채권관리", "예상 회수등급, 예상 처리기간, 유사사건, 채권규모 기반 우선순위 판단을 보조한다."),
        ("4.12 RAG 상담", "상담사례와 법률정보를 구분 검색하고 근거와 전문가 이관 기준을 표시한다."),
    ]:
        h2(doc, title)
        para(doc, body)

    h1(doc, "5. 위험진단 설계")
    add_table(doc, ["구성", "역할", "주의사항"], [
        ["규칙 엔진", "등기상 압류·가압류·경매개시, 근저당, 채권최고액, 계약서 버전 불일치 등 명시 규칙 판단", "법률 효과를 단정하지 않고 검토 필요 표시"],
        ["통계 및 ML 위험유사도", "과거 사고사례군과 입력 계약의 패턴 유사도 산정", "정상계약 기준시점 연결 전에는 사고확률 표현 금지"],
        ["RAG 설명", "상담사례와 법률 안내를 근거로 위험 사유 설명", "법률자문이 아니라 정보 제공"],
        ["위험점수", "규칙 기반 위험요소, 데이터 신뢰도, 모델 산출물을 조합", "점수보다 근거와 조치 우선"],
        ["A/B/C/D 등급", "A 낮음, B 주의, C 고위험, D 진행보류 권고", "등급별 CTA와 전문가 상담 여부 표시"],
        ["누락 처리", "API 실패, 입력 누락, 표본 부족을 별도 상태로 표시", "임의 보간으로 확신도 과장 금지"],
    ], widths=[3.2, 7.0, 6.3])
    add_table(doc, ["위험요소", "가중치 설계 방향", "사용자 설명"], [
        ["압류·가압류·경매개시", "최상위 고위험 규칙", "계약 진행 전 전문가 상담 및 HUG/법률 확인 권고"],
        ["근저당권·채권최고액", "보증금 대비 부담, 말소 여부, 설정일 고려", "말소 예정만으로 해소 처리하지 않고 말소 확인 필요"],
        ["전세가율", "주변 실거래가와 공시가격 보조지표 조합", "공시가격은 시장가격이 아니므로 보조지표"],
        ["선순위보증금", "다가구·다세대에서 수동 입력과 증빙 요구", "전체 임차보증금 미확인 시 Partial data"],
        ["계약서 버전 불일치", "계약 해시와 HUG 제출본 비교", "최종본 일치 전 진행 보류"],
    ], widths=[3.3, 6.0, 7.2])

    h1(doc, "6. 위험요소별 상세 프로토콜")
    protocols = [
        ("근저당", "CODEF 조회로 을구의 근저당권, 채권최고액, 설정일, 말소 여부를 구조화한다. 말소 예정은 EvidenceRequested, 말소 확인 후 Verified로 전환한다."),
        ("압류·가압류", "갑구의 압류·가압류·경매개시 여부를 고위험으로 분류한다. 해소 전 계약 진행 권고를 하지 않는다."),
        ("선순위보증금", "다가구는 전체 선순위보증금을 API만으로 확정하기 어렵다. 임대인 증빙과 사용자 확인을 요구한다."),
        ("국세·지방세", "세금체납 자동조회는 해커톤 MVP에서 구현하지 않는다. 관련 증빙 제출 또는 제도 연계 필요를 표시한다."),
        ("전세가율", "국토부 실거래가, 전월세 실거래가, 공시가격 보조정보를 결합하되 시장가격 확정값으로 표현하지 않는다."),
        ("다운계약·이중계약", "계약서 버전, 보증금, HUG 제출본 해시 불일치와 상담사례를 통해 의심 플래그로만 표시한다."),
        ("계약서 버전 불일치", "초안·보완·최종·HUG 제출본의 해시와 문서버전이 다르면 EvidenceRequested 또는 AtRisk로 전환한다."),
        ("잔금일과 대항력 공백", "전입신고·확정일자·잔금일의 시간차를 사용자 체크리스트로 관리한다."),
        ("보증보험", "HUG 보증 가능 여부는 실제 내부 심사와 다르며, MVP는 입력조건 기반 사전 체크로 한정한다."),
        ("다가구·다세대 구조", "건축물대장으로 용도와 전유부를 확인하되, 선순위보증금은 별도 증빙이 필요하다."),
        ("계약 전 최종 체크", "위험 해소 조건, 증빙 검증, 계약 버전, 조회시각, 데이터 출처, 블록체인 상태를 최종 확인한다."),
    ]
    for name, desc in protocols:
        h2(doc, f"6.{protocols.index((name, desc))+1} {name}")
        para(doc, desc)

    h1(doc, "7. 개인 임대인 트랙")
    add_kv_table(doc, "개인 임대인 트랙", [
        ("입력정보", "임대인 유형, 주소, 보증금, 계약기간, 등기 조회 동의, 필요 시 증빙"),
        ("위험진단", "등기·시세·건축물·계약서 규칙 중심"),
        ("D-90", "반환계획 요청, 응답 여부, 증빙 제출, 미응답 조기경보"),
        ("반환계획", "반환 예정일, 반환 재원, 보완증빙, 임차인 확인"),
        ("사고 후 절차", "미반환 의심 접수, HUG 이행청구 안내, 채권관리 인계"),
        ("한계", "세금체납·실제 지급능력·은행잔고 등은 자동판단하지 않는다."),
    ])

    h1(doc, "8. 법인·개인사업자 트랙")
    h2(doc, "8.1 비식별 법인키 없음 가정")
    para(doc, "기본 설계는 발제사가 비식별 법인키를 제공하지 않는다는 가정으로 작성한다. 이 경우 동일 법인의 여러 물건 직접 연결, 연쇄사고 추적, 법인 단위 그래프 분석, 보증계약→사고→대위변제→경매→배당의 법인 단위 연결은 구현 불가능하다.")
    h2(doc, "8.2 집단 위험분석")
    add_table(doc, ["사용 변수", "분석 방식", "출력 표현"], [[
        "법인/개인사업자 구분, 지역, 주택형태, 사고사유, 사고금액, 대위변제금액, 보증기간, 사고 발생시점",
        "동일 집단의 사고사례군 내 상대적 위험도와 분포 비교",
        "법인사업자·해당 지역·해당 주택유형·해당 보증금 구간의 과거 사고사례군 내 상대적 위험도가 높음"
    ]], widths=[5.5, 5.3, 5.7])
    h2(doc, "8.3 외부 공개정보 활용")
    para(doc, "사용자가 계약서 또는 동의를 통해 실제 사업자등록번호·법인등록번호를 제공하는 경우에만 사업자 상태, OpenDART, 법인 등기, 공개 감사보고서, 공개 재무정보 조회를 확장한다.")
    h2(doc, "8.4 법인 회생·파산 분기")
    para(doc, "OpenDART 주요사항 공시나 공개자료에서 회생·파산 관련 신호를 탐색하되, 모든 비상장 소형 임대법인이 공시 대상은 아니므로 누락 가능성을 표시한다.")
    h2(doc, "8.5 다물건 추적의 현재 한계")
    para(doc, "동일 법인키가 없으면 동일 법인의 다른 계약으로 위험이 확산되는지 실제 그래프 분석을 할 수 없다.")
    h2(doc, "8.6 향후 비식별키 제공 시 확장")
    add_bullets(doc, ["그래프 기반 동일 법인 다물건 연결", "최초 사고 후 후속 사고 발생 간격 분석", "연쇄위험 조기경보", "법인별 사고집중도", "다른 임차인 공동대응 후보 탐지"])

    h1(doc, "9. 사고 후 회수 및 채권관리")
    add_table(doc, ["항목", "설계", "현재 한계"], [
        ["데이터 정의", "사고금액, 대위변제금액, 경매/공매 유형, 배당금, 처리기간, 주택유형, 지역", "공통 키와 회차별 경매 데이터 확인 필요"],
        ["예상 배당률", "HUG 경매·배당데이터 기반 회수등급 또는 예상 배당률 예측", "표본과 변수 부족 시 등급화"],
        ["예상 처리기간", "경공매 신청부터 배당 또는 처리까지의 일수 예측", "절차별 기준일 정의 필요"],
        ["채권 우선순위", "예상 회수등급, 처리기간, 채권규모, 위험근거를 조합", "업무량·정책 기준은 발제사 확인"],
        ["유사사건 분석", "지역·유형·금액·권리관계 유사사례 검색", "개별 법인 확률로 표현 금지"],
        ["What-if", "최적 매각가, 회차별 낙찰률 등은 추가 데이터 필요", "현재 데이터만으로 검증 불가"],
        ["대안 매각채널", "온비드 공매를 별도 모집단으로 분석", "법원 경매와 절차가 다름"],
    ], widths=[3.0, 7.0, 6.5])

    h1(doc, "10. 데이터 설계")
    para(doc, "상세 데이터 수집 절차 및 API 호출 명세는 별도 Markdown 문서로 관리하며, 본 개발설계보고서에서는 데이터의 역할, 시스템 연결 구조 및 구현 우선순위를 정의한다.")
    add_table(doc, ["데이터", "역할", "연결 키", "MVP/확장"], [
        ["발제사 상담데이터", "분쟁유형·진행단계·전문가 이관 모델 및 RAG", "상담ID, 익명화 메타데이터", "MVP/PoC"],
        ["HUG 사고·경매·배당데이터", "회수등급·처리기간 예측", "사고ID, 물건ID, 절차일자", "PoC"],
        ["등기부 조회결과", "계약 전 권리관계 위험진단", "주소 정규화 키, 조회시각", "MVP"],
        ["실거래가", "전세가율·이상계약 탐지", "법정동코드, 주택유형, 기간", "MVP"],
        ["건축물대장", "주택유형·용도·면적 확인", "건물관리번호", "MVP"],
        ["법인 공개정보", "법인 리스크 보조", "법인명/사업자번호/고유번호", "확장"],
    ], widths=[3.6, 6.0, 4.0, 2.9])
    add_bullets(doc, [
        "키 부재 대응: 동일 법인키가 없을 때는 집단 분석으로 한정한다.",
        "파생변수: 전세가율, 채권최고액/보증금 비율, 조회시점 경과일, D-90 응답 지연일, 사고사례군 유사도.",
        "개인정보: 원문 문서와 상세 식별자는 오브젝트 스토리지에 암호화 저장하고 온체인에는 해시만 기록한다.",
        "상세 API 호출 코드와 파라미터는 별도 Markdown에서 관리한다.",
    ])

    h1(doc, "11. Open API 및 외부 연계 설계")
    p = para(doc, "CODEF 부동산 등기부등본 API 공식 제품 페이지: ")
    add_hyperlink(p, "https://developer.codef.io/products/public/each/ck/real-estate-register", "https://developer.codef.io/products/public/each/ck/real-estate-register")
    h2(doc, "11.1 CODEF 등기부 API")
    add_flow(doc, "CODEF 권장 조회 흐름", [
        "사용자 주소 입력 -> 도로명주소 정규화 -> CODEF 주소 검색 -> 주소 후보 선택",
        "-> 건물·동·호 확정 -> 등기부등본 조회 -> 표제부·갑구·을구 핵심 항목 구조화",
        "-> 위험진단 -> 조회결과 및 문서 해시 기록",
    ])
    para(doc, "주소로 조회를 시작할 수 있으나 모든 입력 주소가 즉시 하나의 등기부로 확정되는 것은 아니다. 도로명주소 또는 지번주소 정규화, 집합건물 동·호 입력, 주소 후보 선택, API 인증·계정·과금·호출 제한 확인이 필요하다. 문서에서는 ‘실시간 확인’ 대신 ‘조회 시점 기준 최신 등기정보 확인’으로 표현한다.")
    add_table(doc, ["등기부 활용 항목", "위험진단 활용"], [
        ["소재지, 건물 유형, 조회일시", "주소와 계약 물건 일치 확인"],
        ["소유자 일치 여부, 소유권 이전", "임대인 권한 및 최근 권리 변동 확인"],
        ["압류, 가압류, 경매개시", "고위험 즉시 플래그"],
        ["근저당권, 채권최고액, 공동담보 여부, 설정일, 말소 여부", "보증금 대비 부담과 말소 검증"],
    ], widths=[7.5, 9.0])
    h2(doc, "11.2~11.10 API 목록 및 fallback")
    add_table(doc, ["API/데이터명", "제공기관", "활용 목적", "사용자 입력값", "시스템 출력값", "인증", "유료", "MVP", "fallback", "개인정보·법률 주의"], api_rows, widths=[2.4, 1.9, 2.7, 2.2, 2.8, 1.1, 1.2, 1.3, 1.9, 2.5], font_size=6.2)
    add_flow(doc, "공통 API 실패 대비 흐름", [
        "API 자동조회 -> 성공 시 결과 정규화 및 위험진단 -> 실패 시 재시도",
        "-> 계속 실패하면 문서 업로드 또는 샘플 데이터 -> 사용자가 추출결과 확인 -> 위험진단",
    ])
    add_bullets(doc, ["실 API 모드", "Mock API 모드", "사전 저장 샘플 JSON", "정상·근저당·압류/가압류 등기부 샘플", "API 타임아웃 처리", "오류 메시지", "재시도 버튼"])

    h1(doc, "12. ML 모델 설계")
    h2(doc, "12.1 Google Colab 개발환경")
    add_bullets(doc, [
        "개발환경: Google Colab, Python 기반 Jupyter Notebook(.ipynb), Google Drive 마운트",
        "라이브러리: pandas, numpy, scikit-learn, LightGBM, CatBoost, XGBoost, SHAP, sentence-transformers 또는 Ko-SBERT",
        "노트북 구성: 데이터 로딩, EDA, 전처리, 모델 학습, 평가, 모델 저장을 셀 단위로 분리",
        "모델 아티팩트: joblib, pickle, booster 파일 등으로 내보내고 FastAPI가 로드해 추론",
        "모델별로 별도 .ipynb 파일을 작성한다.",
    ])
    add_table(doc, ["모델", "목적", "입력 변수", "타깃 변수", "후보", "평가지표", "저장/연동"], [
        ["모델 1 상담 분류", "아이엔 상담데이터 기반 분쟁유형·진행단계 분류 및 전문가 이관 필요도", "상담본문 임베딩, 상담 메타데이터, 진행 상태", "분쟁유형, 진행단계, 이관 필요 여부", "Ko-SBERT+LogReg, CatBoost, LightGBM", "F1, macro-F1, confusion matrix", "joblib; FastAPI /ml/counsel-classify"],
        ["모델 2 회수등급", "HUG 경매·배당데이터 기반 예상 배당률 또는 회수등급 예측", "지역, 주택유형, 권리관계, 사고금액, 대위변제금액, 경매 특성", "배당률 또는 회수등급", "LightGBM, CatBoost, XGBoost", "MAE/RMSE 또는 macro-F1, calibration", "booster/joblib; /ml/recovery-grade"],
        ["모델 3 처리기간", "경공매 신청부터 배당 또는 처리까지 예상 소요기간 예측", "절차 시작일, 물건유형, 지역, 채권규모, 경매/공매 유형", "처리 소요일", "LightGBM Regressor, CatBoost Regressor", "MAE, RMSE, P50/P80 error", "joblib; /ml/expected-days"],
    ], widths=[2.3, 4.0, 3.3, 2.5, 2.4, 2.2, 2.8], font_size=6.8)
    add_table(doc, ["설계 항목", "기준"], [
        ["전처리", "결측치 플래그화, 범주형 인코딩, 금액 로그변환 후보, 날짜 파생변수, 텍스트 임베딩 캐시"],
        ["데이터 분할", "시간 누수 방지를 위해 가능하면 기준일 기반 train/valid/test 분리"],
        ["검증 방식", "K-fold는 보조로 사용하고 운영 시점 재현성을 위해 시간 기반 holdout 우선"],
        ["SHAP", "개별 예측 근거와 전체 feature importance를 UI와 관리자 화면에 연결"],
        ["버전관리", "데이터 버전, 노트북 버전, 모델 파일, 평가 지표, 학습일시, 해시를 model_versions에 저장"],
        ["한계", "정상계약 데이터가 없을 때 사고확률, 실제 부도확률, 실제 미반환확률 표현 금지"],
        ["허용 표현", "과거 사고사례 위험유사도, 고위험 패턴 충족도, 위험등급, 사고사례군 내 상대적 위험도"],
        ["합성데이터", "데모, API·UI 테스트, 개인정보 없는 시연, 클래스 불균형 보조, 원본 표본이 충분할 때 제한적 증강으로만 사용"],
    ], widths=[3.0, 13.5])

    h1(doc, "13. RAG 설계")
    add_table(doc, ["항목", "설계 기준"], [
        ["상담데이터 임베딩", "sentence-transformers 또는 Ko-SBERT로 상담본문·요약·분쟁유형을 임베딩"],
        ["검색 필터", "지역, 주택유형, 임대인 유형, 분쟁단계, 위험요소, 사건상태"],
        ["유사사례", "임차인에게는 익명화 요약만 표시하고 담당자에게는 권한 범위 내 상세 표시"],
        ["법률정보와 상담사례 구분", "법령·제도 정보와 상담 경험칙을 UI에서 분리"],
        ["답변 구조", "요약, 근거, 필요한 증빙, 다음 행동, 전문가 이관 기준"],
        ["근거 표시", "출처, 상담사례 ID, 문서 버전, 조회시각"],
        ["환각 방지", "근거 없는 법률 결론 금지, 모르는 경우 확인 필요로 답변"],
        ["전문가 이관", "압류·경매·고액 근저당·미반환 의심·법률 쟁점 발생 시 이관"],
    ], widths=[4.0, 12.5])

    h1(doc, "14. 블록체인 설계")
    para(doc, "블록체인은 계약·검증·상태변경의 공동 감사장부이자 임대차 생애주기 신뢰 레이어다. 법적 강제집행 기능이나 ML 예측 기능을 수행한다고 주장하지 않는다.")
    add_table(doc, ["역할", "기록 대상", "온체인 저장 기준"], [
        ["계약 버전 공증", "계약 초안, 보완계약, 최종 계약, HUG 제출 계약서", "문서 해시, 버전, 상태, 제출처"],
        ["등기/API 조회결과 공증", "조회시각, 조회결과 해시, 근저당·압류 상태, 검증기관, 위험등급", "상세 개인정보와 금액 원문 제외"],
        ["위험조건 해소 이력", "근저당 존재, 말소 예정, 말소 확인, 위험등급 변경", "상태와 해시"],
        ["기관별 검증 서명", "임대인 제출, 아이엔 검토, CODEF/API 조회, HUG 확인", "verifier role과 event log"],
        ["계약 상태머신", "Draft, Diagnosed, EvidenceRequested, EvidenceSubmitted, Verified, ContractFinalized, D90Requested, ReturnPlanSubmitted, AtRisk, IncidentReported, TransferredToHUG, Closed", "상태변경 이벤트"],
        ["D-90 반환준비", "반환계획 요청, 응답, 증빙, 보완요청, 미제출, 반환곤란", "요약 상태와 해시"],
        ["사고 후 인계", "사고의심 등록, HUG 이행청구, 채권관리 인계, 회수분석 실행", "인계 이벤트와 참조 ID"],
    ], widths=[3.2, 8.0, 5.3])
    add_table(doc, ["온체인 금지", "블록체인이 수행하지 못하는 것"], [[
        "계약서 원문, 등기부 원문, 주민등록번호, 사업자등록번호 원문, 상세주소, 예금·소득정보, 상담본문, 세금자료 원문",
        "허위문서 내용의 진실성 자동 판정, 근저당 말소, 세금체납 조회, 원화 잔금 자동이체, 압류, 경매개시, HUG 대위변제 승인, 법원 배당순위 결정, ML 위험예측"
    ]], widths=[8.2, 8.3])
    add_table(doc, ["스마트컨트랙트 개념", "설명"], [
        ["ContractRecord", "계약 ID 해시, 문서 버전, 현재 상태"],
        ["EvidenceRecord", "증빙 유형, 문서 해시, 제출자 역할, 제출시각"],
        ["RegistrySnapshot", "조회기관, 조회시각, 결과 해시, 위험등급"],
        ["VerificationRecord", "검증자, 검증결과, 서명 또는 상태"],
        ["ContractStatus", "상태머신 enum"],
        ["RiskGradeHistory", "등급 변경 이력"],
        ["D90ReturnPlanStatus", "D-90 요청·응답 상태"],
        ["IncidentTransferRecord", "사고 후 HUG 인계 이력"],
        ["verifier/administrator role", "검증자와 운영자 권한 분리"],
        ["event log", "상태변경, 버전등록, 검증등록 이벤트"],
        ["document version/hash verification", "원문 없이 해시로 일치 확인"],
    ], widths=[5.0, 11.5])
    add_flow(doc, "D-90 트리거 구조", [
        "PostgreSQL 계약 만료일 -> 백엔드 스케줄러 -> D-90 대상 계약 조회 -> D-90 이벤트 생성",
        "-> DB 저장 -> 임대인·임차인 알림 -> 상태 변경 -> 상태변경 해시 또는 이벤트를 블록체인 기록",
    ])
    para(doc, "스케줄러 후보는 APScheduler, Celery Beat, cron, Cloud Scheduler다. 블록체인은 스케줄러가 아니라 상태와 이력을 검증하는 레이어다.")

    h1(doc, "15. 시스템 아키텍처")
    add_flow(doc, "전체 논리 아키텍처", [
        "Next.js/React + TypeScript + Tailwind + shadcn/ui",
        "        | REST/HTTPS",
        "FastAPI + Pydantic + SQLAlchemy + Alembic",
        "        |-- External API Gateway: CODEF, 도로명주소, 국토부, 건축물대장, OpenDART, 온비드",
        "        |-- AI/ML Service: Colab 학습 아티팩트 로드, SHAP 설명",
        "        |-- RAG: sentence-transformers/Ko-SBERT + pgvector 또는 Chroma + LLM API",
        "        |-- Blockchain Adapter: ethers.js, Hardhat, Polygon Amoy, MetaMask",
        "        |-- Scheduler/Notification: APScheduler/Celery Beat + Email/SMS/Kakao/FCM",
        "PostgreSQL + pgvector | Object Storage(Supabase Storage 또는 S3) | Mock JSON",
    ])
    add_table(doc, ["영역", "선택 기술", "선택 이유", "대안"], [
        ["Frontend", "Next.js 또는 React, TypeScript, Tailwind CSS, shadcn/ui", "빠른 화면 구현, 타입 안정성, 컴포넌트 재사용", "Vite React, MUI"],
        ["Backend", "FastAPI, Pydantic, SQLAlchemy, Alembic", "Python ML 연동과 API 개발 속도", "Django, NestJS"],
        ["Database", "PostgreSQL, pgvector, Supabase 또는 로컬 PostgreSQL", "관계형 데이터와 벡터검색 동시 지원", "MongoDB Atlas Vector Search 고려 가능"],
        ["Object Storage", "Supabase Storage 또는 S3 호환", "문서 원문 분리 저장", "로컬 MinIO"],
        ["ML", "Colab, scikit-learn, LightGBM/CatBoost/XGBoost, SHAP", "노트북 기반 실험과 설명가능성", "AutoML은 보조"],
        ["RAG", "sentence-transformers, Ko-SBERT, pgvector/Chroma, LLM API", "한국어 상담사례 검색", "OpenSearch vector"],
        ["Blockchain", "Solidity, Hardhat, ethers.js, Polygon Amoy, MetaMask", "테스트넷 시연과 이벤트 검증", "Sepolia"],
        ["Deployment", "Vercel, Render/Railway, Supabase, 환경변수 Secret", "1인 개발 배포 속도", "Fly.io, AWS"],
    ], widths=[2.2, 4.2, 6.2, 3.9], font_size=7.2)
    para(doc, "사용자가 고려 중인 MongoDB는 문서형 일반 데이터와 벡터 검색을 함께 구성할 수 있다는 장점이 있다. 다만 본 설계의 기본안은 트랜잭션성 계약 상태, 감사로그, SQL 기반 조인이 중요한 점을 고려해 PostgreSQL+pgvector로 둔다. MongoDB Atlas Vector Search는 RAG 저장소 또는 프로토타입 단순화 대안으로 검토한다.")

    h1(doc, "16. 데이터 흐름도")
    flows = {
        "계약 전 진단": "입력 -> 주소 정규화 -> 등기/시세/건축물 조회 -> 규칙·ML 위험유사도 -> 위험등급 -> 보완조건",
        "CODEF 조회": "주소 후보 -> 동·호 확정 -> 등기 조회 -> 갑구/을구 구조화 -> 결과 해시 -> DB/블록체인",
        "위험등급 산정": "규칙 엔진 -> 모델 추론 -> 데이터 신뢰도 -> A/B/C/D -> 사용자 설명",
        "증빙 검증": "요청 -> 제출 -> 문서추출 -> 사용자 확인 -> 검증자 승인 -> 상태변경",
        "블록체인 기록": "문서/결과 해시 -> verifier 서명 -> event log -> tx hash 저장",
        "D-90": "만료일 조회 -> 알림 -> 반환계획 -> 미응답/반환곤란 -> AtRisk",
        "사고 후 채권관리": "사고접수 -> HUG 인계 -> 회수등급/처리기간 -> 우선순위 -> 담당자 액션",
        "ML 학습·배포": "Drive 데이터 -> Colab EDA/전처리/학습/평가 -> artifact 저장 -> FastAPI 로드",
    }
    for k, v in flows.items():
        h2(doc, f"16.{list(flows.keys()).index(k)+1} {k}")
        para(doc, v)

    h1(doc, "17. User Flow")
    add_table(doc, ["사용자 흐름", "단계", "예외"], [
        ["임차인 계약 전", "홈 -> 주소/계약 입력 -> API 조회 -> 위험결과 -> 보완요청", "API 실패 시 업로드 대체"],
        ["임대인 증빙", "알림 -> 요청 확인 -> 파일 제출 -> 검증 상태 확인", "보완요청 또는 반려"],
        ["계약 확정", "버전 비교 -> 최종본 확정 -> 해시 기록 -> 타임라인", "해시 불일치 시 보류"],
        ["계약 중 모니터링", "상태 조회 -> 변경 감지 -> 알림", "주기 재조회 동의·비용 필요"],
        ["D-90", "스케줄러 이벤트 -> 반환계획 요청 -> 응답/미응답 -> 경보", "반환곤란 입력"],
        ["사고 후 임차인", "미반환 의심 -> 사고 접수 -> HUG 안내 -> 인계", "중복 접수 방지"],
        ["HUG 담당자", "대시보드 -> 사건 상세 -> 유사사건/회수예측 -> 액션", "모델 실패 시 규칙·수동 우선"],
        ["법인 임대인", "집단 위험분석 -> 공개정보 조회(동의 시) -> 상담/검증", "비식별 법인키 없으면 직접 연결 금지"],
        ["API 실패 예외", "재시도 -> 업로드 -> 추출 확인 -> 위험진단", "샘플 데이터 모드"],
    ], widths=[3.2, 9.0, 4.3])

    h1(doc, "18. IA 및 내비게이션 구조")
    add_table(doc, ["영역", "Sitemap", "권한"], [
        ["임차인", "홈, 계약 입력, API 조회, 위험결과, 위험상세, 계약 타임라인, D-90, 사고 대응", "tenant"],
        ["임대인", "증빙 요청, 증빙 제출, 검증 진행, 반환계획 제출", "landlord"],
        ["HUG 관리자", "대시보드, 사건 상세, 유사사건 비교, 법인 위험분석, 담당자 액션", "hug_admin"],
        ["아이엔", "상담 RAG, 증빙 검토, 전문가 이관, 위험해설", "advisor"],
        ["공통", "로그인, 알림, 권한 오류, 블록체인 이력, 문서 업로드", "authorized"],
    ], widths=[3.0, 10.0, 3.5])

    h1(doc, "19. UI/UX 상세 설계")
    add_bullets(doc, [
        "디자인 원칙: 점수보다 조치, 출처와 조회시각 표시, 해결 가능 위험과 불가능 위험 분리, 개인정보 최소 노출.",
        "위험등급 컬러: A는 녹색, B는 노랑, C는 주황, D는 빨강을 사용하되 색상만으로 의미를 전달하지 않는다.",
        "Empty·Loading·Error 상태는 모든 API 조회 화면에서 별도 컴포넌트로 설계한다.",
        "모바일에서는 계약 입력과 위험 결과를 단계형 카드로 분리하고 긴 표는 요약 후 상세 접기로 제공한다.",
        "접근성: 등급 배지는 텍스트 라벨을 병기하고 주요 CTA는 키보드 포커스와 명확한 오류 메시지를 제공한다.",
    ])
    add_table(doc, ["위험등급 화면 표시항목", "설명"], [
        ["위험등급", "A/B/C/D 및 계약 진행 권고"],
        ["위험요소", "근저당, 압류, 선순위보증금, 전세가율, 계약서 버전 등"],
        ["위험의 근거", "조회 항목, 규칙, 모델 설명, SHAP 주요 변수"],
        ["해결할 수 있는 위험", "말소 확인, 특약 반영, 증빙 제출 등"],
        ["해결할 수 없는 위험", "경매개시, 권리관계 중대 하자 등"],
        ["필요한 증빙", "등기부, 말소확인, 계약서, 반환계획 등"],
        ["전문가 상담 여부", "필요/권고/선택"],
        ["데이터 조회시각·출처", "CODEF, 국토부, 건축물대장 등"],
        ["모델 버전", "model_versions 기준"],
        ["블록체인 검증 여부", "pending/confirmed와 tx hash"],
    ], widths=[5.0, 11.5])

    h1(doc, "20. Wireframe")
    for sid, name, user, purpose, entry, req, api, display, comp, cta, states, nxt, auth, err, mobile in screen_rows:
        h2(doc, f"20.{screen_rows.index([sid, name, user, purpose, entry, req, api, display, comp, cta, states, nxt, auth, err, mobile])+1} {name}")
        add_kv_table(doc, f"{sid} {name} 화면 정의", [
            ("화면 ID", sid), ("화면명", name), ("사용자", user), ("화면 목적", purpose), ("진입 경로", entry),
            ("필수 입력값", req), ("조회 API", api), ("화면에 표시할 데이터", display), ("핵심 컴포넌트", comp),
            ("주요 CTA", cta), ("상태", states + "; Default, Loading, Success, Empty, Partial data, API failure, Validation error, Permission denied, Blockchain pending, Blockchain confirmed, Model inference failure 중 해당 상태 적용"),
            ("다음 화면", nxt), ("권한", auth), ("오류 상황", err), ("모바일 고려사항", mobile)
        ])

    h1(doc, "21. API 명세")
    add_table(doc, ["영역", "Endpoint", "Method", "설명"], [
        ["인증", "/api/auth/login", "POST", "역할 기반 로그인"],
        ["계약", "/api/contracts", "POST/GET", "계약 생성·목록"],
        ["위험진단", "/api/contracts/{id}/diagnose", "POST", "규칙·ML 위험진단 실행"],
        ["등기조회", "/api/registry/search, /api/registry/snapshot", "POST", "주소 후보와 등기 스냅샷"],
        ["증빙", "/api/evidence-requests, /api/evidences", "POST/GET", "증빙 요청·제출"],
        ["타임라인", "/api/contracts/{id}/timeline", "GET", "상태 변경 이력"],
        ["D-90", "/api/contracts/{id}/return-plan", "POST/GET", "반환계획 요청·응답"],
        ["사고", "/api/incidents", "POST/GET", "사고 접수·조회"],
        ["회수예측", "/api/incidents/{id}/recovery-prediction", "POST", "회수등급·처리기간 추론"],
        ["관리자", "/api/admin/dashboard", "GET", "HUG 대시보드"],
        ["블록체인", "/api/blockchain/records", "POST/GET", "해시 기록과 tx 조회"],
        ["외부 API Adapter", "/api/adapters/{provider}", "POST", "외부 API 호출 래핑"],
        ["오류코드", "API_TIMEOUT, PROVIDER_AUTH_FAILED, PARTIAL_DATA, MODEL_FAILED, CHAIN_PENDING", "-", "공통 오류"],
    ], widths=[2.6, 5.6, 2.0, 6.3], font_size=7.5)

    h1(doc, "22. 데이터베이스 설계")
    add_table(doc, ["테이블", "주요 필드", "설명"], db_rows, widths=[3.6, 7.0, 5.9], font_size=7.0)

    h1(doc, "23. 해커톤 시연 시나리오")
    add_table(doc, ["시나리오", "시연 내용", "설명 포인트"], [
        ["계약 전 위험발견", "주소 입력 후 등기·시세 조회와 위험등급 표시", "점수가 아니라 근거와 조치 중심"],
        ["근저당 말소 보완", "근저당 발견 -> 말소 증빙 요청 -> 말소 확인", "위험 해소 검증"],
        ["API 조회 결과 공증", "조회결과 해시와 tx hash 표시", "원문 개인정보 온체인 금지"],
        ["계약서 버전 공증", "초안/최종/HUG 제출본 해시 비교", "버전 불일치 방지"],
        ["D-90 미응답 조기경보", "스케줄러 이벤트와 알림 상태", "스마트컨트랙트 자동실행 아님"],
        ["사고 후 채권관리", "사고 인계 후 회수등급·처리기간", "업무 우선순위 보조"],
        ["법인 집단위험", "법인사업자·지역·유형별 사고사례군 비교", "개별 사고확률 표현 금지"],
        ["API 실패 fallback", "CODEF 실패 -> PDF 업로드 -> 추출 확인", "시연 안정성"],
    ], widths=[3.4, 7.0, 6.1])

    h1(doc, "24. 개발 로드맵")
    add_table(doc, ["순서", "작업", "산출물"], [
        ["1", "데이터/API 수집 Markdown 작성", "API 목록, 키 필요 여부, 샘플 JSON"],
        ["2", "백엔드 기본 스키마와 Mock API", "FastAPI, PostgreSQL, api_call_logs"],
        ["3", "CODEF/주소/실거래가 Adapter", "실 API 또는 Mock 전환"],
        ["4", "프론트엔드 핵심 화면", "TEN-01~TEN-05, HUG-01"],
        ["5", "ML Notebook 개발", "상담 분류, 회수등급, 처리기간 .ipynb"],
        ["6", "RAG 검색", "임베딩, 필터, 답변 템플릿"],
        ["7", "블록체인 컨트랙트", "Hardhat, Solidity, Amoy 배포"],
        ["8", "D-90 스케줄러와 알림", "APScheduler/Celery Beat"],
        ["9", "통합·예외·데모 리허설", "샘플 데이터와 발표 시나리오"],
    ], widths=[1.5, 7.0, 8.0])

    h1(doc, "25. 테스트 계획")
    add_table(doc, ["테스트", "기준"], [
        ["기능 테스트", "계약 생성, 주소 후보, 위험진단, 증빙요청, 타임라인"],
        ["API 테스트", "성공, timeout, 인증실패, rate limit, fallback"],
        ["모델 테스트", "입력 스키마, 결측, 추론 실패, 버전 표시, SHAP 생성"],
        ["스마트컨트랙트 테스트", "상태전이, 권한, 이벤트, 해시 검증, 중복 기록"],
        ["예외 테스트", "Partial data, permission denied, validation error"],
        ["개인정보 테스트", "온체인 원문 금지, 로그 마스킹, 다운로드 권한"],
        ["데모 리허설", "실 API 모드와 Mock 모드 전환"],
        ["API 장애 대비", "샘플 JSON, 재시도 버튼, 오류 메시지"],
    ], widths=[4.0, 12.5])

    h1(doc, "26. 보안·개인정보·법률")
    add_bullets(doc, [
        "개인정보 최소수집: 진단에 필요한 항목만 수집하고 해시·마스킹·권한 분리를 적용한다.",
        "온체인 개인정보 금지: 원문 문서, 주민등록번호, 사업자등록번호 원문, 상세주소, 상담본문을 저장하지 않는다.",
        "해시의 한계: 해시는 특정 문서와의 일치 확인 도구이며 문서 내용의 진실성을 보장하지 않는다.",
        "API 키 관리: 환경변수와 Secret Manager를 사용하고 프론트엔드에 노출하지 않는다.",
        "로그 관리: API 요청/응답 원문 저장을 제한하고 장애 분석에 필요한 최소 메타데이터만 보관한다.",
        "법률적 한계: 서비스 안내는 법률자문이 아니라 정보 제공이며, 고위험 건은 전문가 상담으로 이관한다.",
        "금융·은행 연계 한계: 실제 원화 이체는 은행 제휴, 본인인증, 금융망 연계, 에스크로 계좌, 취소·분쟁 처리, 전자금융·법률 검토가 필요하다.",
    ])

    h1(doc, "27. 리스크 및 한계")
    add_table(doc, ["리스크", "영향", "대응방안"], [
        ["데이터 키 부재", "계약·물건·사고 연결 제한", "공통 키 제공 여부 질문, 집단 분석으로 대체"],
        ["정상계약 부재", "실제 사고확률 모델 불가", "위험유사도·상대위험도 표현"],
        ["API 비용·승인", "실 API 시연 불안정", "Mock API와 샘플 JSON"],
        ["CODEF 조회 실패", "등기 자동분석 중단", "PDF 업로드와 사용자 확인"],
        ["법인정보 누락", "비상장 법인 분석 제한", "OpenDART 누락 가능성 표시"],
        ["경매 회차 데이터 부족", "최적 매각가/처리기간 모델 제한", "추가 데이터 요청"],
        ["임대인 참여 유인", "증빙 제출 지연", "요청 사유와 계약 진행 CTA 명확화"],
        ["모델 일반화", "표본 편향", "시간기반 검증과 SHAP, 과대해석 금지"],
        ["법적 집행 한계", "블록체인으로 권리 실행 불가", "감사장부 역할로 제한"],
    ], widths=[3.5, 6.0, 7.0])

    h1(doc, "28. 발제사 확인 질문")
    add_bullets(doc, [
        "비식별 법인키 제공 가능 여부",
        "정상계약 데이터 제공 여부와 기준시점",
        "공통 계약·물건·사고키 존재 여부",
        "경매 회차별 데이터와 배당 데이터 범위",
        "HUG 기존 등기 모니터링 또는 알림 업무 존재 여부",
        "D-90 반환계획 확인 제도화 가능성",
        "채권관리 우선순위 기준과 현재 업무량·처리비용",
        "블록체인 감사장부에 대한 실무 수요",
        "외부 API 및 Mock 사용 허용 범위",
        "CODEF 실제 계정, 과금, 호출 제한, 테스트 가능 여부",
    ])

    h1(doc, "29. 평가기준 대응")
    add_table(doc, ["평가기준", "대응 논리"], [
        ["독창성", "점수 앱이 아니라 위험 해소 검증과 사고 전후 연결을 구현"],
        ["적합성", "HUG·아이엔의 상담, 보증, 채권관리 흐름에 맞춘 생애주기 구조"],
        ["충실성", "API, ML, RAG, 블록체인, UI, DB, 테스트까지 구현 기준 제시"],
        ["활용성", "임차인·임대인·상담·HUG 담당자별 화면과 업무 액션 정의"],
        ["기술경쟁력", "CODEF, 공공데이터, Colab ML, RAG, Amoy 블록체인 통합"],
        ["데이터 활용성", "정상계약 부재와 비식별키 부재를 명확히 반영"],
        ["확장성", "비식별 법인키, 은행 에스크로, 기관 PoC로 확장 가능"],
    ], widths=[3.0, 13.5])

    h1(doc, "30. 최종 결론")
    para(doc, "HUG 안심전세 체인은 전세 위험을 단순 점수로 환산하는 서비스가 아니라, 위험 발견, 위험 해소 조건 제안, 증빙 검증, 계약 및 조회 이력 공증, 만료 전 반환 준비 확인, 사고 후 HUG 채권관리 인계를 하나의 흐름으로 연결하는 주거안전 인프라다. MVP는 1인 개발자가 구현 가능한 범위에서 외부 API와 Mock을 병행하고, ML은 실제 사고확률이 아닌 위험유사도·회수등급·처리기간 보조 모델로 한정한다. 블록체인은 법적 집행 장치가 아니라 계약·검증·상태변경의 공동 감사장부로 위치시킨다.")

    h1(doc, "부록 A. 기능 우선순위표")
    add_table(doc, ["기능", "우선순위", "구분"], [
        ["주소 정규화", "P0", "MVP"], ["CODEF 등기조회/업로드 대체", "P0", "MVP"], ["위험등급과 보완요청", "P0", "MVP"], ["계약 버전 해시", "P0", "MVP"], ["D-90 스케줄러", "P1", "MVP/데모"], ["회수등급 모델", "P1", "PoC"], ["법인 그래프", "P3", "확장"], ["은행 에스크로", "P3", "확장"]
    ], widths=[8.0, 3.0, 5.5])
    h1(doc, "부록 B. MVP·Mock·PoC·확장 구분표")
    add_table(doc, ["구분", "기능"], [["MVP", "핵심 입력, 등기/API 조회, 위험진단, 보완요청, 해시 기록, 관리자 Mock"], ["Mock", "샘플 등기부, 샘플 실거래가, 샘플 HUG 사건"], ["PoC", "실 HUG 데이터와 CODEF 계정 연계"], ["확장", "비식별 법인키 그래프, 은행 에스크로, 기관 자동연계"]], widths=[3.0, 13.5])
    h1(doc, "부록 C. 데이터/API 준비 목록")
    add_table(doc, ["준비물", "상태"], [["CODEF 계정·과금·호출 테스트", "확인 필요"], ["공공데이터포털 키", "확인 필요"], ["샘플 JSON", "준비 필요"], ["정상/근저당/압류 등기부 샘플", "준비 필요"], ["HUG 사고·경매·배당 데이터", "발제사 확인 필요"]], widths=[8.0, 8.5])
    h1(doc, "부록 D. 화면 정의서 요약")
    add_table(doc, ["화면 ID", "화면명", "사용자", "핵심 CTA"], [[r[0], r[1], r[2], r[9]] for r in screen_rows], widths=[2.2, 5.0, 3.0, 6.3], font_size=7.5)
    h1(doc, "부록 E. 용어 정의")
    add_table(doc, ["용어", "정의"], [["위험유사도", "과거 사고사례군과 입력 계약의 패턴 유사 정도"], ["D-90", "계약 만료 90일 전 반환계획 확인 이벤트"], ["온체인", "블록체인에 기록되는 해시·상태·이벤트"], ["오프체인", "DB와 스토리지에 저장되는 원문·상세 데이터"], ["Verifier", "검증 상태를 기록할 권한을 가진 주체"]], widths=[4.0, 12.5])
    h1(doc, "부록 F. 스마트컨트랙트 상태 목록")
    add_bullets(doc, ["Draft", "Diagnosed", "EvidenceRequested", "EvidenceSubmitted", "Verified", "ContractFinalized", "D90Requested", "ReturnPlanSubmitted", "AtRisk", "IncidentReported", "TransferredToHUG", "Closed"])
    h1(doc, "부록 G. ML 산출물 목록")
    add_bullets(doc, ["model1_counsel_classification.ipynb", "model2_recovery_grade.ipynb", "model3_expected_duration.ipynb", "preprocess_pipeline.joblib", "model artifact", "metrics.json", "shap_summary.json", "model_card.md"])
    h1(doc, "부록 H. 발표용 핵심 메시지")
    add_bullets(doc, [
        "위험을 점수화하는 앱이 아니라 위험 해소를 증명하는 플랫폼",
        "계약 전 진단, 계약 이행 검증, 만료 전 조기개입, 사고 후 채권관리를 하나로 연결",
        "블록체인은 원문 저장소가 아니라 공동 감사장부",
        "정상계약 데이터 부재 시 사고확률이 아니라 위험유사도로 표현",
        "해커톤 MVP는 실 API와 Mock을 병행해 안정적으로 시연",
    ])

    doc.core_properties.title = "HUG × 아이엔 안심주거 생태계 개발기획 및 상세설계보고서"
    doc.core_properties.author = "Codex"
    doc.core_properties.subject = "DIVE 2026 개발설계보고서"
    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
